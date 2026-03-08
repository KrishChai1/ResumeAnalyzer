#!/usr/bin/env python3
"""
╔══════════════════════════════════════════════════════════════════════════════╗
║         RESUME PARSER API v8.0 - COMPLETE AGENTIC FRAMEWORK                   ║
╠══════════════════════════════════════════════════════════════════════════════╣
║                                                                               ║
║  AGENTIC ARCHITECTURE:                                                        ║
║  =====================                                                        ║
║                                                                               ║
║  ┌────────────────────────────────────────────────────────────────────────┐  ║
║  │                        ORCHESTRATION LAYER                              │  ║
║  │  ┌─────────────┐  ┌─────────────┐  ┌─────────────┐  ┌─────────────┐   │  ║
║  │  │ EXTRACTION  │─▶│ VALIDATION  │─▶│ AI ENHANCE  │─▶│ OUTPUT      │   │  ║
║  │  │ AGENT       │  │ AGENT       │  │ AGENT       │  │ AGENT       │   │  ║
║  │  │ (10 Pattern)│  │ (Scoring)   │  │ (Claude)    │  │ (JSON)      │   │  ║
║  │  └─────────────┘  └─────────────┘  └─────────────┘  └─────────────┘   │  ║
║  └────────────────────────────────────────────────────────────────────────┘  ║
║                                                                               ║
║  v8.0 IMPROVEMENTS:                                                           ║
║  ==================                                                           ║
║  • Enhanced name extraction with filename fallback                            ║
║  • 10 experience extraction patterns                                          ║
║  • Role-based skill inference (Java Developer = Java experience)              ║
║  • Short date parsing (Jul24, Feb'20, May'14)                                 ║
║  • Two-column PDF handling                                                    ║
║  • Improved validation with lenient scoring                                   ║
║  • Better responsibility extraction                                           ║
║                                                                               ║
║  SUPPORTED FORMATS: PDF, DOCX (tables/textboxes), TXT, ZIP                    ║
║                                                                               ║
╚══════════════════════════════════════════════════════════════════════════════╝

Version: 7.0.0 (Production)
"""

import os
import io
import re
import json
import zipfile
import asyncio
from typing import Optional, Dict, List, Any, Tuple, Set
from datetime import datetime
from dataclasses import dataclass, field
from enum import Enum

# Third-party imports
try:
    from dateutil.relativedelta import relativedelta
except ImportError:
    relativedelta = None

from fastapi import FastAPI, HTTPException, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field

# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                           CONFIGURATION                                       ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

VERSION = "8.8.0"
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")

# Month name to number mapping
MONTH_MAP = {
    'jan': 1, 'january': 1, 'feb': 2, 'february': 2, 'mar': 3, 'march': 3,
    'apr': 4, 'april': 4, 'may': 5, 'jun': 6, 'june': 6, 'jul': 7, 'july': 7,
    'aug': 8, 'august': 8, 'sep': 9, 'sept': 9, 'september': 9,
    'oct': 10, 'october': 10, 'nov': 11, 'november': 11, 'dec': 12, 'december': 12
}

# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                         SKILL CATEGORIES                                      ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

SKILL_CATEGORIES = {
    "Programming Languages": [
        "python", "java", "javascript", "typescript", "c++", "c#", "go", "golang",
        "rust", "ruby", "php", "swift", "kotlin", "scala", "cobol", "bash", "shell",
        "powershell", "sql", "plsql", "t-sql", "r", "matlab", "perl", "groovy",
        "vb.net", "asp.net", "vbscript"
    ],
    "Web & Frameworks": [
        "html", "css", "react", "angular", "vue", "node.js", "nodejs", "express",
        "django", "flask", "fastapi", ".net", "asp.net", "spring", "spring boot",
        "hibernate", "struts", "ejb", "jquery", "bootstrap", "tailwind", "spring batch",
        "spring mvc", "spring jpa", "webflux", "micronaut", "quarkus"
    ],
    "Cloud Platforms": [
        "aws", "azure", "gcp", "google cloud", "ibm cloud", "openstack",
        "ec2", "s3", "lambda", "ecs", "ecr", "eks", "rds", "vpc", "route 53",
        "cloudformation", "cloudwatch", "sns", "sqs", "dynamodb", "aurora",
        "azure devops", "azure sql", "bigquery", "dataproc", "dataflow", "composer"
    ],
    "Data Engineering": [
        "etl", "data warehouse", "data pipeline", "informatica", "ssis", "ssas", "ssrs",
        "talend", "apache airflow", "airflow", "apache spark", "spark", "pyspark",
        "apache kafka", "kafka", "snowflake", "databricks", "dbt", "glue", "redshift",
        "data lake", "bigquery", "nifi", "hdfs", "oozie", "sqoop", "hue"
    ],
    "Databases": [
        "mongodb", "cassandra", "redis", "postgresql", "mysql", "oracle", "db2",
        "sql server", "nosql", "dynamodb", "elasticsearch", "neo4j", "teradata",
        "vertica", "singlestore", "hive", "hbase", "couchdb", "mariadb", "postgres"
    ],
    "DevOps & CI/CD": [
        "docker", "kubernetes", "k8s", "jenkins", "gitlab", "github actions",
        "ci/cd", "cicd", "concourse", "helm", "gitops", "devops", "devsecops",
        "azure devops", "octopus", "bamboo", "travis ci", "circleci", "argo",
        "maven", "gradle", "ant"
    ],
    "Infrastructure & Virtualization": [
        "terraform", "ansible", "puppet", "chef", "vagrant", "packer",
        "vmware", "vmware vsphere", "vcloud", "virtualization", "hyper-v",
        "openshift", "rancher", "istio", "consul", "vault", "cloudformation"
    ],
    "Testing & QA": [
        "selenium", "pytest", "unittest", "testng", "junit", "cucumber", "behave",
        "bdd", "tdd", "manual testing", "automation testing", "regression testing",
        "api testing", "postman", "functional testing", "uat", "quality assurance",
        "jmeter", "sonarqube", "loadrunner", "cypress", "playwright", "mockito"
    ],
    "Security": [
        "qualys", "crowdstrike", "mend", "uptycs", "snyk", "veracode",
        "siem", "splunk", "palo alto", "firewall", "iam", "pam", "mfa",
        "cisco security", "checkpoint", "fortinet", "zscaler", "oauth", "jwt"
    ],
    "Business Intelligence & Visualization": [
        "tableau", "power bi", "grafana", "prometheus", "elk stack", "kibana",
        "ssrs", "ssas", "obiee", "looker", "qlik", "spotfire", "datadog",
        "elastic search", "logstash", "business objects"
    ],
    "Project Management & Collaboration": [
        "jira", "confluence", "servicenow", "agile", "scrum", "kanban",
        "waterfall", "pmp", "prince2", "itil", "ms project", "clarity-ppm",
        "monday.com", "asana", "trello", "azure boards", "smartsheet"
    ],
    "Messaging & Integration": [
        "kafka", "rabbitmq", "ibm mq", "activemq", "sqs", "sns", "pubsub",
        "rest", "soap", "graphql", "grpc", "api gateway", "mulesoft", "apigee",
        "ibm message broker"
    ],
    "Version Control": [
        "git", "github", "gitlab", "bitbucket", "svn", "mercurial", "perforce",
        "tfs", "vsts"
    ],
    "Operating Systems": [
        "linux", "ubuntu", "rhel", "centos", "debian", "unix", "windows server",
        "macos", "redhat", "fedora", "alpine", "aix", "solaris"
    ],
    "Business Domains": [
        "banking", "finance", "financial services", "insurance", "healthcare",
        "telecom", "retail", "e-commerce", "manufacturing", "pharma",
        "investment banking", "capital markets", "payments", "billing"
    ],
    "SAP & ERP": [
        "sap", "sap s/4 hana", "sap hana", "sap sd", "sap mm", "sap fi", "sap co",
        "sap wm", "sap ewm", "sap tm", "sap cs", "sap pm", "sap pp", "sap qm",
        "sap bw", "sap abap", "sap fico", "sap basis", "sap brim", "sap rar",
        "sap otc", "sap gts", "oracle erp", "oracle ebs", "peoplesoft",
        "dynamics 365", "d365", "netsuite", "workday", "sap btp", "sap c4c"
    ]
}

# Flatten all skills for quick lookup
ALL_SKILLS: Set[str] = set()
for skills in SKILL_CATEGORIES.values():
    ALL_SKILLS.update(s.lower() for s in skills)

# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                    ROLE-BASED SKILL INFERENCE (v8.0)                         ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

ROLE_SKILL_MAP = {
    "java": ["java", "spring", "spring boot", "maven", "junit", "hibernate", "rest"],
    "python": ["python", "pytest", "django", "flask", "pandas"],
    "devops": ["docker", "kubernetes", "jenkins", "ci/cd", "ansible", "terraform", "git"],
    "data engineer": ["sql", "etl", "spark", "kafka", "airflow", "snowflake", "python"],
    "frontend": ["javascript", "react", "angular", "html", "css", "typescript"],
    "backend": ["java", "python", "sql", "rest", "microservices", "spring"],
    "full stack": ["java", "javascript", "sql", "react", "angular", "spring", "node.js"],
    "cloud": ["aws", "gcp", "azure", "docker", "kubernetes", "terraform"],
    "etl": ["sql", "etl", "informatica", "data pipeline", "ssis", "python"],
    "scrum master": ["agile", "scrum", "jira", "confluence"],
    "project manager": ["agile", "jira", "confluence", "scrum", "project management"],
    "qa": ["selenium", "junit", "testing", "automation testing", "jmeter"],
    "test": ["selenium", "junit", "testing", "automation testing", "pytest"],
    "architect": ["microservices", "aws", "docker", "kubernetes", "spring"],
    "analyst": ["sql", "excel", "tableau", "power bi", "data analysis"],
    "database": ["sql", "oracle", "mysql", "postgresql", "mongodb"],
    "security": ["security", "firewall", "iam", "oauth", "ssl"],
}

# Location patterns
LOCATION_CITIES = [
    'Philadelphia', 'Chicago', 'Dallas', 'Plano', 'Houston', 'Atlanta', 
    'Bangalore', 'Bengaluru', 'Hyderabad', 'North Chicago', 'New York',
    'San Francisco', 'Des Moines', 'Seattle', 'Boston', 'Denver', 'Austin',
    'Los Angeles', 'Portland', 'Charlotte', 'Raleigh', 'Nashville', 'Tampa',
    'Noida', 'Gurgaon', 'Pune', 'Mumbai', 'Chennai', 'Kolkata', 'France',
    'North America', 'India', 'USA', 'UK', 'Germany', 'Singapore', 'Canada'
]
LOCATION_PATTERN = '|'.join(re.escape(city) for city in LOCATION_CITIES[:35])

# Tech terms that should NOT be treated as names
TECH_TERMS = {
    'SQL', 'ETL', 'GCP', 'AWS', 'API', 'XML', 'JSON', 'HTML', 'CSS', 'SSIS',
    'Spring', 'JPA', 'Java', 'Hibernate', 'EJB', 'REST', 'SOAP', 'Kafka',
    'Docker', 'Maven', 'Gradle', 'Jenkins', 'Git', 'Linux', 'Unix', 'Python',
    'React', 'Angular', 'Node', 'Vue', 'MongoDB', 'Redis', 'MySQL', 'Oracle',
    'Spark', 'Hadoop', 'Hive', 'Scala', 'Kibana', 'Elastic', 'Grafana', 'Prometheus',
    'DOMAIN', 'SKILLS', 'JIRA', 'BI', 'IT', 'JAVA', 'PYTHON', 'TECHNICAL',
    'AND', 'OR', 'THE', 'FOR', 'WITH'
}

# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                           DATA CLASSES                                        ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

@dataclass
class ValidationResult:
    """Result from Validation Agent."""
    score: int = 100
    issues: List[str] = field(default_factory=list)
    fixes: List[str] = field(default_factory=list)
    missing_fields: List[str] = field(default_factory=list)
    needs_ai_enhancement: bool = False


class ParseTextRequest(BaseModel):
    text: str = Field(..., min_length=50, description="Resume text to parse")
    use_ai_validation: bool = Field(default=True, description="Enable AI enhancement")
    filename: Optional[str] = Field(default=None, description="Original filename")


class ExtractSkillsRequest(BaseModel):
    text: str = Field(..., min_length=10, description="Text to extract skills from")


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                          FASTAPI APPLICATION                                  ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

app = FastAPI(
    title="Resume Parser API",
    description=f"""
## Enterprise Resume Parser v{VERSION} - Complete Agentic Framework

### 🏗️ Architecture
Multi-agent system with specialized components:

| Agent | Purpose | Details |
|-------|---------|---------|
| **Extraction Agent** | Pattern-based parsing | 10 regex patterns for various formats |
| **Validation Agent** | Quality assurance | Scoring (0-100), auto-fixes |
| **AI Enhancement Agent** | Gap filling | Claude API for missing fields |
| **Output Agent** | Standardization | Clean JSON generation |

### 📄 v8.0 Improvements
- Enhanced name extraction with filename fallback
- Role-based skill inference (Java Developer = Java experience)
- Short date parsing (Jul24, Feb'20)
- Two-column PDF handling
- Better responsibility extraction

### 📁 Supported File Formats
- **PDF**: Standard and multi-column layouts
- **DOCX**: Tables, text boxes, complex formatting
- **TXT**: Plain text
- **ZIP**: Archives containing text files
    """,
    version=VERSION,
    docs_url="/docs",
    redoc_url="/redoc"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                       FILE EXTRACTION UTILITIES                               ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

def detect_file_type(content: bytes, filename: str) -> str:
    """Detect actual file type from magic bytes."""
    if content[:4] == b'PK\x03\x04':
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                names = zf.namelist()
                if any('word/document.xml' in n for n in names):
                    return 'docx'
                elif any(n.endswith('.txt') for n in names):
                    return 'zip_text'
                return 'zip'
        except:
            pass
    elif content[:5] == b'%PDF-':
        return 'pdf'
    elif content[:4] == b'\xd0\xcf\x11\xe0':
        return 'doc'
    
    # Check if it's actually text content
    try:
        text_sample = content[:1000].decode('utf-8', errors='strict')
        if text_sample.startswith(('#', '##', '**', 'PROFESSIONAL', 'SUMMARY', 'RESUME', '-')):
            return 'txt'
        printable_ratio = sum(1 for c in text_sample if c.isprintable() or c in '\n\r\t') / len(text_sample)
        if printable_ratio > 0.9:
            return 'txt'
    except:
        pass
    
    ext = filename.lower().split('.')[-1] if filename else ''
    return ext if ext in ['pdf', 'docx', 'doc', 'txt'] else 'unknown'


def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF using multiple methods."""
    text = ""
    
    # Try pdfplumber first
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            return text
    except ImportError:
        pass
    except Exception:
        pass
    
    # Try PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(content))
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception:
        pass
    
    # Try zipfile method for special PDFs
    if not text.strip():
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as z:
                text_files = sorted(
                    [n for n in z.namelist() if n.endswith('.txt')],
                    key=lambda x: int(x.split('.')[0]) if x.split('.')[0].isdigit() else 999
                )
                for tf in text_files:
                    text += z.read(tf).decode('utf-8', errors='ignore') + "\n"
        except:
            pass
    
    return text


def extract_text_from_doc(content: bytes) -> str:
    """Extract text from legacy .doc files using pure Python (v8.6.3).
    
    Uses olefile if available, falls back to direct binary extraction.
    """
    text_parts = []
    
    # Method 1: Try olefile for OLE compound documents
    olefile_success = False
    try:
        import olefile
        if olefile.isOleFile(io.BytesIO(content)):
            ole = olefile.OleFileIO(io.BytesIO(content))
            if ole.exists('WordDocument'):
                word_data = ole.openstream('WordDocument').read()
                
                # Extract ASCII text from binary
                current_text = ""
                for byte in word_data:
                    if 32 <= byte <= 126:  # Printable ASCII
                        current_text += chr(byte)
                    elif byte in [10, 13]:  # Newlines
                        if current_text.strip():
                            text_parts.append(current_text.strip())
                        current_text = ""
                    else:
                        # Non-printable: add space to preserve word boundaries
                        if current_text and not current_text.endswith(' '):
                            current_text += ' '
                
                if current_text.strip():
                    text_parts.append(current_text.strip())
                olefile_success = True
            ole.close()
    except ImportError:
        pass  # olefile not installed
    except Exception:
        pass  # olefile failed
    
    # Method 2: Direct binary extraction (always runs as fallback)
    if not olefile_success or not text_parts:
        text_parts = []  # Reset
        current_text = ""
        for byte in content:
            if 32 <= byte <= 126:  # Printable ASCII
                current_text += chr(byte)
            elif byte in [10, 13]:  # Newlines
                if len(current_text.strip()) > 3:
                    text_parts.append(current_text.strip())
                current_text = ""
            else:
                # Non-printable: add space
                if current_text and not current_text.endswith(' '):
                    current_text += ' '
        
        if len(current_text.strip()) > 3:
            text_parts.append(current_text.strip())
    
    # Filter: keep only meaningful text lines
    filtered_parts = []
    for part in text_parts:
        # Normalize whitespace
        part = ' '.join(part.split())
        if len(part) < 5:
            continue
        
        # Always keep lines with contact info or dates
        if '@' in part:
            # Extract the meaningful part (might have garbage prefix)
            email_match = re.search(r'([A-Za-z][A-Za-z\s]*)?\s*Email[:\s]*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+)', part, re.IGNORECASE)
            if email_match:
                # Extract name if present before Email
                name_part = email_match.group(1)
                if name_part and len(name_part.strip()) > 2:
                    # Clean up - take only the name part
                    name_clean = name_part.strip()
                    # Find where the real name starts (after garbage)
                    words = name_clean.split()
                    good_words = []
                    for w in words:
                        if len(w) > 2 and w[0].isupper() and w[1:].islower():
                            good_words.append(w)
                    if good_words:
                        filtered_parts.append(' '.join(good_words) + ' Email: ' + email_match.group(2))
                        continue
            filtered_parts.append(part)
            continue
        if re.search(r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}', part):
            filtered_parts.append(part)
            continue
        if re.search(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{4}', part, re.IGNORECASE):
            filtered_parts.append(part)
            continue
        
        # Check text quality - ratio of readable chars
        readable = sum(1 for c in part if c.isalnum() or c.isspace() or c in '.,;:-@()&/')
        if readable / len(part) < 0.6:
            continue
        
        # Check average word length (skip garbled text)
        words = [w for w in part.split() if len(w) > 0]
        if not words:
            continue
        avg_len = sum(len(w) for w in words) / len(words)
        if avg_len < 2:
            continue
        
        # Skip lines that are mostly single characters
        single_char_words = sum(1 for w in words if len(w) == 1)
        if len(words) > 3 and single_char_words / len(words) > 0.5:
            continue
        
        filtered_parts.append(part)
    
    # Deduplicate
    seen = set()
    result = []
    for part in filtered_parts:
        if part not in seen:
            seen.add(part)
            result.append(part)
    
    return '\n'.join(result)


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX including tables and text boxes (v8.6.0)."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        
        all_text = []
        
        # First check for name in first few paragraphs (common location)
        for para in doc.paragraphs[:5]:
            text = para.text.strip()
            if text:
                all_text.append(text)
        
        # Then extract tables (contain name/header for some resumes)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    all_text.append(' | '.join(row_text))
        
        # Then remaining paragraphs
        for para in doc.paragraphs[5:]:
            text = para.text.strip()
            if text:
                all_text.append(text)
        
        # Try to extract text boxes
        try:
            xml_str = doc.element.xml
            pattern = r'<w:txbxContent[^>]*>(.*?)</w:txbxContent>'
            matches = re.findall(pattern, xml_str, re.DOTALL)
            for match in matches:
                text_pattern = r'<w:t[^>]*>([^<]+)</w:t>'
                texts = re.findall(text_pattern, match)
                if texts:
                    combined = ' '.join(texts)
                    combined = ' '.join(combined.split())
                    if combined and len(combined) > 5:
                        all_text.append(combined)
        except:
            pass
        
        return '\n'.join(all_text)
    except ImportError:
        # Fallback: try reading as plain text
        try:
            return content.decode('utf-8', errors='ignore')
        except:
            return ""
    except Exception:
        # Fallback: try reading as plain text
        try:
            return content.decode('utf-8', errors='ignore')
        except:
            return ""


def extract_text_from_zip(content: bytes) -> str:
    """Extract text from ZIP archive."""
    text_parts = []
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            for name in sorted(zf.namelist()):
                if name.endswith('.txt'):
                    text_parts.append(zf.read(name).decode('utf-8', errors='ignore'))
    except:
        pass
    return '\n'.join(text_parts)


def extract_text_intelligent(content: bytes, filename: str) -> str:
    """Intelligently extract text based on actual file type."""
    file_type = detect_file_type(content, filename)
    
    if file_type == 'pdf':
        return extract_text_from_pdf(content)
    elif file_type == 'doc':
        return extract_text_from_doc(content)
    elif file_type == 'docx':
        return extract_text_from_docx(content)
    elif file_type == 'zip_text':
        return extract_text_from_zip(content)
    elif file_type == 'zip':
        text = extract_text_from_zip(content)
        if not text.strip():
            text = extract_text_from_docx(content)
        return text
    else:
        try:
            return content.decode('utf-8', errors='ignore')
        except:
            return ""


def normalize_text(text: str) -> str:
    """Clean and normalize resume text."""
    if not text:
        return ""
    
    replacements = {
        'â€"': '–', 'â€™': "'", 'â€œ': '"', 'â€': '"',
        'Ã©': 'é', 'Ã¨': 'è', 'Ã ': 'à',
        '\u2013': '-', '\u2014': '-', '–': '-',
        '\u2019': "'", '\u2018': "'",
        '\u201c': '"', '\u201d': '"',
        '\u2022': '•', '\u00a0': ' ',
        '\r\n': '\n', '\r': '\n', '\t': ' '
    }
    
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    text = re.sub(r'Pres\s*ent', 'Present', text, flags=re.IGNORECASE)
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                        DATE PARSING UTILITIES (v8.0)                         ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

def parse_date(text: str) -> Tuple[Optional[int], Optional[int], bool]:
    """Parse date string to (year, month, is_present)."""
    if not text:
        return None, None, False
    
    text = text.strip().lower()
    
    if any(p in text for p in ['present', 'current', 'now', 'ongoing', 'till date']):
        now = datetime.now()
        return now.year, now.month, True
    
    # Month Year format (full year)
    match = re.search(
        r'(jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|'
        r'aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)'
        r'\s*[,.]?\s*(\d{4})', text
    )
    if match:
        month = MONTH_MAP.get(match.group(1)[:3])
        year = int(match.group(2))
        return year, month, False
    
    # Compact with optional apostrophe: Jan2021, Jul24, Feb'20, May'14
    match = re.search(r"(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)['\s]?(\d{2,4})", text, re.IGNORECASE)
    if match:
        month = MONTH_MAP.get(match.group(1).lower())
        year_str = match.group(2)
        year = 2000 + int(year_str) if len(year_str) == 2 and int(year_str) < 50 else (
            1900 + int(year_str) if len(year_str) == 2 else int(year_str)
        )
        return year, month, False
    
    # Month with space and 2-digit year: "July 15"
    match = re.search(r'(jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|'
                      r'aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)'
                      r'\s+(\d{2})(?:\s|$)', text)
    if match:
        month = MONTH_MAP.get(match.group(1)[:3])
        year_str = match.group(2)
        year = 2000 + int(year_str) if int(year_str) < 50 else 1900 + int(year_str)
        return year, month, False
    
    # Just year
    match = re.search(r'\b(19\d{2}|20\d{2})\b', text)
    if match:
        return int(match.group(1)), None, False
    
    return None, None, False


def parse_short_date(date_str: str) -> Tuple[Optional[int], Optional[int], bool]:
    """Parse short date formats like Jul24, Jun21, Feb'20 (v8.0)."""
    if not date_str:
        return None, None, False
    
    date_str = date_str.strip().lower().replace("'", "")
    
    if date_str in ['present', 'current', 'now']:
        now = datetime.now()
        return now.year, now.month, True
    
    # MonYY format (Jul24 -> July 2024)
    m = re.match(r'([a-z]+)(\d{2})$', date_str)
    if m:
        month_name = m.group(1)[:3]
        year_short = int(m.group(2))
        year = 2000 + year_short if year_short < 50 else 1900 + year_short
        if month_name in MONTH_MAP:
            return year, MONTH_MAP[month_name], False
    
    # Fall back to standard parse
    return parse_date(date_str)


def calculate_duration(start_year: int, start_month: int, end_year: int, end_month: int) -> int:
    """Calculate duration in months (inclusive)."""
    if relativedelta:
        start_dt = datetime(start_year, start_month or 1, 1)
        end_dt = datetime(end_year, end_month or 12, 1)
        delta = relativedelta(end_dt, start_dt)
        return max(1, delta.years * 12 + delta.months + 1)
    else:
        # Simple calculation without dateutil
        return max(1, (end_year - start_year) * 12 + (end_month or 12) - (start_month or 1) + 1)


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                     EXTRACTION AGENT - CONTACT INFO                          ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

def extract_contact(text: str) -> Dict[str, str]:
    """Extract contact information from resume text."""
    contact = {'email': '', 'phone': '', 'linkedin': '', 'location': ''}
    
    # Email
    match = re.search(r'\b([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\b', text)
    if match:
        contact['email'] = match.group(1).lower()
    
    # Phone - updated patterns for various formats
    phone_patterns = [
        r'(?:Mob|Phone|Tel|Mobile|Cell|Ph|Contact)[\s]*(?:No)?[:\s]*(\+?[\d\s\-().]{10,})',
        r'(\+1\s*\(?\d{3}\)?[-.\s]*\d{3}[-.\s]*\d{4})',
        r'(\(\d{3}\)[-.\s]*\d{3}[-.\s]*\d{4})',
        r'(\+\d{1,3}[-.\s]*\(?\d{3}\)?[-.\s]*\d{3}[-.\s]*\d{4})',
        r'(\d{3}[-.\s]?\d{3}[-.\s]?\d{4})',
    ]
    for pattern in phone_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            phone = re.sub(r'[^\d+\-() ]', '', match.group(1)).strip()
            if len(re.sub(r'\D', '', phone)) >= 10:
                contact['phone'] = phone
                break
    
    # LinkedIn
    linkedin_patterns = [
        r'linkedin\.com/in/([\w-]+)',
        r'LinkedIn[:\s]+([\w-]+)',
    ]
    for pattern in linkedin_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            username = match.group(1)
            if username.lower() not in ['summary', 'profile', 'in', 'phone', 'email']:
                contact['linkedin'] = f"www.linkedin.com/in/{username}"
                break
    
    # Location - prioritize first 10 lines (header area)
    header_text = '\n'.join(text.split('\n')[:10])
    
    # Common US state abbreviations and countries
    state_country_pattern = r'(AL|AK|AZ|AR|CA|CO|CT|DE|FL|GA|HI|ID|IL|IN|IA|KS|KY|LA|ME|MD|MA|MI|MN|MS|MO|MT|NE|NV|NH|NJ|NM|NY|NC|ND|OH|OK|OR|PA|RI|SC|SD|TN|TX|UT|VT|VA|WA|WV|WI|WY|India|USA|UK|Canada|Germany|Australia|Singapore|UAE|Karnataka|Maharashtra|Tamil Nadu|Telangana|Andhra Pradesh)'
    
    # First try to find location in header (first 10 lines)
    # Pattern: City, STATE or City STATE
    header_loc_match = re.search(rf'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)[,\s]+{state_country_pattern}\b', header_text)
    if header_loc_match:
        city = header_loc_match.group(1).strip()
        state = header_loc_match.group(2).strip()
        contact['location'] = f"{city}, {state}"
    else:
        # Fallback: search full text but only first occurrence
        for city in LOCATION_CITIES[:30]:
            pattern = rf'\b{re.escape(city)}[,\s]+{state_country_pattern}\b'
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                contact['location'] = f"{city}, {match.group(1)}"
                break
    
    return contact


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                 EXTRACTION AGENT - NAME EXTRACTION (v8.0)                    ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

def extract_name(text: str, filename: str = "") -> Tuple[str, str, str]:
    """
    Extract first, middle, last name from resume (v8.0).
    Enhanced with filename fallback and better tech term filtering.
    """
    lines = [l for l in text.split('\n')][:150]
    
    skip_patterns = [
        'resume', 'cv', 'curriculum vitae', 'key expertise', 'professional experience',
        'education', 'skills', 'summary', 'objective', 'technical', 'profile',
        'results-driven', 'experienced', 'skilled', 'dedicated', 'total',
        'data engineering', 'core competencies', 'with over', 'contact', 'houston',
        'years of experience', 'developer with', 'professional summary', 'spring',
        'java', 'python', 'aws', 'gcp', 'sql', 'micro', 'kafka', 'rest', 'xml',
        'docker', 'kubernetes', 'jenkins', 'git', 'linux', 'cloud', 'database',
        'web flux', 'hibernate', 'junit', 'jboss', 'tomcat', 'maven', 'gradle',
        'layout table', 'ph:', 'email:', 'phone:', 'cell:', 'domain', 'governance'
    ]
    
    def split_name_parts(parts: List[str]) -> Tuple[str, str, str]:
        """Split name parts into first, middle, last."""
        if len(parts) == 2:
            return parts[0], "", parts[1]
        elif len(parts) == 3:
            return parts[0], parts[1], parts[2]
        elif len(parts) >= 4:
            return parts[0], ' '.join(parts[1:-1]), parts[-1]
        return "", "", ""
    
    # Strategy 0a: Check if first non-empty line is a single capitalized name followed by Email line
    # This handles formats like "Venkata\nEmail: ..." or "Ramaswamy Tati                Email: ..."
    for i, line in enumerate(lines[:10]):
        clean_line = line.strip()
        if not clean_line:
            continue
        
        # Check if this line contains a name followed by Email on same line or next line
        # Pattern: "Name                       Email: ..."
        name_email_match = re.match(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s+(?:Email|Mobile|Phone|Contact)', clean_line)
        if name_email_match:
            name_str = name_email_match.group(1).strip()
            parts = name_str.split()
            if 1 <= len(parts) <= 4:
                if not any(p.upper() in TECH_TERMS for p in parts):
                    if len(parts) == 1:
                        return parts[0], "", ""
                    return split_name_parts(parts)
        
        # Check if line is just a single capitalized name and next line has Email/Contact
        if '@' not in clean_line and ':' not in clean_line:
            parts = clean_line.split()
            if 1 <= len(parts) <= 4 and all(p.isalpha() and p[0].isupper() for p in parts):
                # Check next line for Email/Contact
                if i + 1 < len(lines):
                    next_line = lines[i + 1].lower().strip()
                    if any(kw in next_line for kw in ['email', 'phone', 'mobile', 'contact', '@']):
                        if not any(skip in clean_line.lower() for skip in skip_patterns):
                            if len(parts) == 1:
                                return parts[0], "", ""
                            return split_name_parts(parts)
    
    # Strategy 0b: Look for name in first 15 lines (handles indented names)
    for line in lines[:15]:
        clean_line = line.strip()
        if not clean_line or clean_line.startswith(('-', ':', '|')):
            continue
        if any(clean_line.lower().startswith(skip) for skip in skip_patterns):
            continue
        if '@' in clean_line or clean_line.endswith(':'):
            continue
        if '|' in clean_line:
            continue
        
        parts = clean_line.split()
        if 2 <= len(parts) <= 3:
            if all(p.isalpha() and p[0].isupper() and (p[1:].islower() if len(p) > 1 else True) for p in parts):
                if not any(p.upper() in TECH_TERMS for p in parts):
                    return split_name_parts(parts)
    
    # Strategy 1: ## Name pattern (markdown header)
    for line in lines[:25]:
        clean_line = line.strip()
        m = re.match(r'^#+\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)\s+Email:', clean_line)
        if m:
            parts = m.group(1).strip().split()
            if 2 <= len(parts) <= 4:
                return split_name_parts(parts)
        m = re.match(r'^#+\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)\s*$', clean_line)
        if m:
            parts = m.group(1).strip().split()
            if 2 <= len(parts) <= 4:
                return split_name_parts(parts)
    
    # Strategy 2: **Bold Name** with optional contact info
    for line in lines[:25]:
        clean_line = line.strip()
        m = re.match(r'^\*\*([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)+)(?:\s+Contact:|[*\s]*$)', clean_line)
        if m:
            name_str = m.group(1).strip()
            name_str = re.sub(r'\s+(Contact|Phone|Email|Cell|Mobile).*$', '', name_str, flags=re.IGNORECASE)
            parts = name_str.split()
            if 2 <= len(parts) <= 4:
                if not any(p.upper() in TECH_TERMS for p in parts):
                    return split_name_parts([p.title() if p.isupper() else p for p in parts])
    
    # Strategy 3: ALL-CAPS name line (PDF sidebars)
    for line in lines:  # Search entire text
        clean_line = ' '.join(line.split()).strip('\r')
        if clean_line.isupper() and 5 < len(clean_line) < 40:
            parts = clean_line.split()
            if 2 <= len(parts) <= 4:
                if not any(p in TECH_TERMS for p in parts):
                    if parts[0] not in ['TECHNICAL', 'PROFESSIONAL', 'WORK', 'EDUCATION', 
                                        'SKILLS', 'CORE', 'KEY', 'DOMAIN', 'SUMMARY']:
                        return split_name_parts([p.title() for p in parts])
    
    # Strategy 4: Standard capitalized name (first 40 lines)
    for line in lines[:40]:
        clean_line = ' '.join(line.split())
        
        if any(clean_line.lower().startswith(skip) for skip in skip_patterns):
            continue
        if clean_line.endswith(':') or '@' in clean_line:
            continue
        if re.match(r'^[\d\s\-+()]+$', clean_line):
            continue
        if ',' in clean_line and not re.match(r'^[A-Z][a-z]+\s*,', clean_line):
            continue
        if len(clean_line) > 50:
            continue
        if clean_line.startswith('>') or clean_line.startswith('-') or clean_line.startswith('*'):
            continue
        if '■' in clean_line or '•' in clean_line or '|' in clean_line:
            continue
        
        name = re.sub(r'\s*[\|].*$', '', clean_line)
        name = re.sub(r'\s*Email:.*$', '', name, flags=re.IGNORECASE)
        name = re.sub(r'\s*Phone:.*$', '', name, flags=re.IGNORECASE)
        name = re.sub(r'\s*Contact:.*$', '', name, flags=re.IGNORECASE)
        name = re.sub(r'\s*,.*$', '', name)
        name = re.sub(r'^#+\s*', '', name)
        name = re.sub(r'^\*+|\*+$', '', name)
        
        parts = name.split()
        if 2 <= len(parts) <= 4:
            if all(p[0].isupper() for p in parts if p):
                if not any(p.upper() in TECH_TERMS or p in TECH_TERMS for p in parts):
                    return split_name_parts(parts)
    
    # Strategy 5: Fallback to filename
    if filename:
        base = re.sub(r'\.(pdf|docx|doc)$', '', filename, flags=re.IGNORECASE)
        base = re.sub(r'[-_]?(Resume|CV)[-_]?\d*$', '', base, flags=re.IGNORECASE)
        base = re.sub(r'[-_]?\d+$', '', base)
        base = re.sub(r'[-_]?(OTC|TM|RAR|GCP|ETL)[-_]?', '', base, flags=re.IGNORECASE)
        parts = re.split(r'[-_]', base)
        non_name_parts = {'resume', 'cv', 'data', 'engineer', 'developer', 'manager', 
                         'consultant', 'sr', 'junior', 'senior', 'lead', 'gcp', 'aws',
                         'java', 'python', 'etl', 'devops', 'it', 'pm', 'scrum', 'master',
                         'project', 'snowflake', 'cloud', 'otc', 'tm', 'rar', 'sap', 'hana'}
        parts = [p.strip() for p in parts if p.strip().lower() not in non_name_parts and len(p) > 1]
        if len(parts) == 1 and len(parts[0]) > 3:
            # Single name from filename - might be first name only (e.g., "Venkata")
            return parts[0].title(), "", ""
        if 2 <= len(parts) <= 4:
            return split_name_parts([p.title() for p in parts])
    
    return "", "", ""


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║              EXTRACTION AGENT - EXPERIENCE (10 PATTERNS) v8.0                ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

def extract_responsibilities(lines: List[str], start_idx: int, end_idx: int) -> List[str]:
    """Extract responsibilities from lines (v8.6 - handles plain text and bullets)."""
    responsibilities = []
    current_resp = ""
    bullet_only_line = False
    
    # Check if the line before start_idx is a responsibility header
    after_resp_header = False
    if start_idx > 0:
        prev_line = lines[start_idx - 1].strip()
        if re.match(r'^(Contribution|Responsibilities)[\s&]*:?\s*$', prev_line, re.IGNORECASE):
            after_resp_header = True
    
    for i in range(start_idx, min(end_idx, len(lines))):
        line = lines[i]
        stripped = line.strip().rstrip('\r')
        
        # Skip empty lines
        if not stripped:
            continue
        
        # Skip "Contribution & Responsibilities:" or similar headers
        if re.match(r'^(Contribution|Responsibilities)[\s&]*:?\s*$', stripped, re.IGNORECASE):
            after_resp_header = True
            continue
        
        # Stop at next Project/Client/Duration section
        if re.match(r'^(Project|Client|Duration|Environment|Description)\s*[:\t]', stripped, re.IGNORECASE):
            break
        
        # Plain text responsibility (sentence starting with action verb)
        if re.match(r'^(Participate|Collaborate|Wrote|Test|Revise|Improve|Involve|Serve|Design|Develop|Manage|Create|Implement|Lead|Support|Analyze|Build|Configure|Maintain|Monitor|Deploy|Review|Ensure|Establish|Coordinate|Execute|Perform|Provide|Work|Assist|Conduct|Document|Deliver|Define|Evaluate|Facilitate|Guide|Handle|Identify|Install|Integrate|Investigate|Launch|Migrate|Optimize|Organize|Oversee|Plan|Prepare|Process|Produce|Program|Research|Resolve|Schedule|Secure|Setup|Streamline|Supervise|Train|Troubleshoot|Update|Upgrade|Validate|Verify|Write)', stripped, re.IGNORECASE):
            if len(stripped) > 20:
                responsibilities.append(stripped[:500])
                after_resp_header = True  # Mark that we're in responsibilities section
                continue
        
        # Stop at section headers
        if re.match(r'^(EDUCATION|TECHNICAL|SKILLS|CERTIFICATIONS|PERSONAL|EXPERIENCE\s*$)', stripped, re.IGNORECASE):
            break
        
        # Stop at next experience (pipe format)
        if re.match(r'^[A-Za-z][^|]*\|[^|]*\|.*\d{4}', stripped):
            break
        
        # Stop at next experience (date pattern)
        if re.search(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}\s*[-–]', stripped, re.IGNORECASE):
            break
        if re.match(r'^\d{1,2}/\d{4}\s*[-–]', stripped):
            break
        
        # Check for standalone bullet
        if stripped in ('■', '•', '-', '*', '>'):
            if current_resp and len(current_resp) > 20:
                responsibilities.append(current_resp.strip()[:500])
            current_resp = ""
            bullet_only_line = True
            continue
        
        # New bullet point with text
        if stripped.startswith(('■', '•', '-', '*', '>')) or re.match(r'^\d+\.', stripped):
            if current_resp and len(current_resp) > 20:
                responsibilities.append(current_resp.strip()[:500])
            current_resp = re.sub(r'^[■•\-\*>\d.]+\s*', '', stripped)
            bullet_only_line = False
            after_resp_header = True
        # Plain text line after "Responsibilities:" header (Steven format)
        elif after_resp_header and stripped and len(stripped) > 30:
            if current_resp and len(current_resp) > 20:
                responsibilities.append(current_resp.strip()[:500])
            current_resp = stripped
        # Text following a standalone bullet
        elif bullet_only_line and stripped:
            current_resp = stripped
            bullet_only_line = False
        # Continuation
        elif current_resp and stripped:
            if line.startswith('    ') or line.startswith('\t') or not stripped[0].isupper() or len(stripped) < 60:
                current_resp += ' ' + stripped
        # Blank line - save current
        elif not stripped and current_resp:
            if len(current_resp) > 20:
                responsibilities.append(current_resp.strip()[:500])
            current_resp = ""
            bullet_only_line = False
    
    if current_resp and len(current_resp) > 20:
        responsibilities.append(current_resp.strip()[:500])
    
    return responsibilities[:15]


def extract_tools(lines: List[str]) -> List[str]:
    """Extract tools/technologies mentioned in lines."""
    tools = []
    text = ' '.join(lines).lower()
    
    for skill in ALL_SKILLS:
        if re.search(rf'\b{re.escape(skill)}\b', text):
            tools.append(skill)
    
    return tools[:20]


def extract_experiences(text: str) -> List[Dict]:
    """
    ╔════════════════════════════════════════════════════════════════════════════╗
    ║                    EXTRACTION AGENT - 10 PATTERNS (v8.0)                    ║
    ╠════════════════════════════════════════════════════════════════════════════╣
    ║ Pattern 1: Standard "Company – Title Date" format                           ║
    ║ Pattern 2: "Worked as X in Y from A to B" format (Madhuri)                  ║
    ║ Pattern 3: Table format "Client: X | Duration: Y" (Ramaswamy)               ║
    ║ Pattern 4: "Title (DateRange) – Client – Employer" (Nageswara)              ║
    ║ Pattern 5: "Company Date" then "Title Location" next line (Khaliq)          ║
    ║ Pattern 6: Pipe format "Title | Date" with company above (Jimmy)            ║
    ║ Pattern 7: "ROLE:" / "DESIGNATION:" keyword format (Sarwer)                 ║
    ║ Pattern 8: "**Client: Company -- Location Date**" (Naveen)                  ║
    ║ Pattern 9: "Title Company - Location" then "MM/YYYY - Present" (Javvaji)    ║
    ║ Pattern 10: "## Title | Company, Location | Date" (Steven markdown)         ║
    ╚════════════════════════════════════════════════════════════════════════════╝
    """
    experiences = []
    text = normalize_text(text)
    lines = text.split('\n')
    
    # Date range pattern
    date_pattern = (
        r'((?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|'
        r'Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4})'
        r'\s*[-–to]+\s*'
        r'((?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|'
        r'Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4}'
        r'|Present|Current|Till\s+[Dd]ate)'
    )
    
    title_keywords = ['Engineer', 'Developer', 'Manager', 'Analyst', 'Lead', 
                      'Consultant', 'Architect', 'Specialist', 'Director',
                      'Admin', 'Coordinator', 'Scrum', 'Master', 'Tester',
                      'Designer', 'Administrator', 'Executive', 'Officer']
    
    # =========================================================================
    # PATTERN 9: MM/YYYY format (Javvaji two-column PDFs)
    # =========================================================================
    mm_yyyy_pattern = r'^(\d{1,2}/\d{4})\s*[-–]\s*(\d{1,2}/\d{4}|Present|Current)$'
    
    for i, line in enumerate(lines):
        clean_line = line.strip()
        date_match = re.match(mm_yyyy_pattern, clean_line, re.IGNORECASE)
        
        if date_match and i > 0:
            start_str = date_match.group(1)
            end_str = date_match.group(2)
            
            start_parts = start_str.split('/')
            if len(start_parts) == 2:
                start_month = int(start_parts[0])
                start_year = int(start_parts[1])
            else:
                continue
            
            if 'present' in end_str.lower() or 'current' in end_str.lower():
                end_year = datetime.now().year
                end_month = datetime.now().month
                is_present = True
            else:
                end_parts = end_str.split('/')
                if len(end_parts) == 2:
                    end_month = int(end_parts[0])
                    end_year = int(end_parts[1])
                    is_present = False
                else:
                    continue
            
            prev_line = lines[i - 1].strip()
            if not prev_line or prev_line.startswith(('•', '-', '*', '■')):
                continue
            
            title, employer, location = "", "", ""
            
            loc_match = re.search(r'\s*[-–]\s*(USA|France|North America|India|UK|Germany|[A-Z][a-z]+(?:\s*&\s*[A-Z][a-z]+)?)$', prev_line, re.IGNORECASE)
            if loc_match:
                location = loc_match.group(1).strip()
                title_company = prev_line[:loc_match.start()].strip()
            else:
                title_company = prev_line
            
            title_end_idx = 0
            for kw in title_keywords:
                kw_lower = kw.lower()
                tc_lower = title_company.lower()
                if kw_lower in tc_lower:
                    idx = tc_lower.find(kw_lower) + len(kw)
                    if idx > title_end_idx:
                        title_end_idx = idx
            
            if title_end_idx > 0:
                title = title_company[:title_end_idx].strip()
                employer = title_company[title_end_idx:].strip()
                employer = re.sub(r'^[\s\-–]+', '', employer).strip()
            else:
                title = title_company
                employer = None
            
            if start_year and end_year:
                duration = calculate_duration(start_year, start_month, end_year, end_month)
                start_date = f"{start_year}-{start_month:02d}"
                end_date = f"{end_year}-{end_month:02d}"
                
                responsibilities = extract_responsibilities(lines, i + 1, i + 40)
                
                is_dup = any(
                    e.get('start_date') == start_date and
                    e.get('title', '').lower()[:20] == title.lower()[:20]
                    for e in experiences
                )
                
                if not is_dup and title:
                    experiences.append({
                        'employer': employer if employer else None,
                        'title': title,
                        'start_date': start_date,
                        'end_date': end_date,
                        'duration_months': duration,
                        'responsibilities': responsibilities,
                        'location': location or None,
                        'pattern': 9
                    })
    
    # =========================================================================
    # PATTERN 4: "Title (DateRange) – Client – Employer" (Nageswara)
    # =========================================================================
    title_date_client_pattern = (
        r'^\s*>?\s*\*?\*?\s*'
        r'([\w\s/]+?)'
        r'\s*\((\w{3,9}\d{2,4})\s*[-–]+\s*(\w{3,9}\d{2,4}|Present|Current)\)'
        r'\s*[-–]+\s*\[?'
        r'([^-–\[\]]+?)'
        r'\s*[-–]+\s*'
        r'\*?\*?\[?\*?\*?'
        r'(.+)$'
    )
    
    for i, line in enumerate(lines):
        clean_line = line.strip()
        match = re.match(title_date_client_pattern, clean_line, re.IGNORECASE)
        if match:
            title = match.group(1).strip()
            start_str = match.group(2)
            end_str = match.group(3)
            client = match.group(4).strip().strip('[').strip(']').strip('*').strip()
            employer = match.group(5).strip().strip(']').strip('[').strip('*').strip()
            
            employer = re.sub(r'\s*\[.*$', '', employer)
            employer = employer.rstrip('_').strip()
            
            start_year, start_month, _ = parse_date(start_str)
            end_year, end_month, is_present = parse_date(end_str)
            
            if start_year and end_year:
                duration = calculate_duration(start_year, start_month or 1, end_year, end_month or 12)
                start_date = f"{start_year}-{(start_month or 1):02d}"
                end_date = f"{end_year}-{(end_month or 12):02d}" if not is_present else f"{datetime.now().year}-{datetime.now().month:02d}"
                
                # Extract responsibilities from table format
                responsibilities = []
                for j in range(i + 1, min(len(lines), i + 40)):
                    resp_line = lines[j].strip()
                    if re.match(r'^\s*>?\s*\*?\*?\s*[\w\s/]+\s*\(\w{3,9}\d{2,4}', resp_line, re.IGNORECASE):
                        break
                    if '|' in resp_line:
                        resp_match = re.findall(r'\|\s*[-•]?\s*([^|]+)', resp_line)
                        for resp in resp_match:
                            resp = resp.strip().strip('-').strip('*').strip()
                            if len(resp) > 25 and not resp.lower().startswith('environment'):
                                responsibilities.append(resp[:500])
                
                is_dup = any(
                    e.get('start_date') == start_date and
                    (e.get('title', '').lower() == title.lower())
                    for e in experiences
                )
                
                if not is_dup:
                    experiences.append({
                        'employer': f"{employer} (Client: {client})" if client != employer else employer,
                        'title': title,
                        'start_date': start_date,
                        'end_date': end_date,
                        'duration_months': duration,
                        'responsibilities': responsibilities[:12],
                        'location': None,
                        'client': client,
                        'pattern': 4
                    })
    
    # =========================================================================
    # PATTERN 8: "**Client: Company -- Location Date**" (Naveen)
    # Handles multiple formats:
    # - Client: Ascent Global Logistics -- Atlanta, GA. Jul24 to Present
    # - Client: Hiscox Inc - Atlanta, Georgia Feb'20 to May'21
    # - Client: UnitedHealth Group, India May'14 to Dec'14 (comma format)
    # - Client: Berkley Technology Services - Des Moines, IA Aug'15 to\nJan'20 (multiline)
    # =========================================================================
    
    i = 0
    while i < len(lines):
        line = lines[i]
        clean_line = re.sub(r'\*\*', '', line.strip())
        
        if not re.match(r'^Client:', clean_line, re.IGNORECASE):
            i += 1
            continue
        
        # Combine with next line in case date spans multiple lines
        combined = clean_line
        if i + 1 < len(lines):
            next_line = re.sub(r'\*\*', '', lines[i + 1].strip())
            if next_line and not next_line.startswith(('Client:', '•', '-', '*', '■')):
                combined += ' ' + next_line
        
        # Try multiple patterns
        match = None
        
        # Pattern A: "Client: Company -- Location. Date to Date"
        # Pattern B: "Client: Company - Location Date to Date"  
        pattern_a = r'Client:\s*(.+?)\s+[-–]+\s+([A-Za-z][A-Za-z\s,]+?)\.?\s+((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z\']*\s*\d{2,4})\s+to\s+((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z\']*\s*\d{2,4}|Present|Current)'
        match = re.match(pattern_a, combined, re.IGNORECASE)
        
        if not match:
            # Pattern B: "Client: Company, Location Date to Date" (comma format)
            pattern_b = r'Client:\s*(.+?),\s*([A-Za-z][A-Za-z\s]+?)\s+((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z\']*\s*\d{2,4})\s+to\s+((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z\']*\s*\d{2,4}|Present|Current)'
            match = re.match(pattern_b, combined, re.IGNORECASE)
        
        if match:
            employer = match.group(1).strip()
            location = match.group(2).strip()
            start_str = match.group(3)
            end_str = match.group(4)
            
            start_year, start_month, _ = parse_date(start_str)
            end_year, end_month, is_present = parse_date(end_str)
            
            if start_year and end_year:
                # Look for title in next 5 lines (skip blank lines)
                title = ""
                title_line_idx = i + 1
                for j in range(i + 1, min(i + 6, len(lines))):
                    next_line = re.sub(r'\*\*', '', lines[j].strip())
                    if not next_line:
                        continue
                    if next_line.startswith(('Client:', 'Responsibilities', 'Description', 'Tech Stack')):
                        break
                    # Skip the date continuation line
                    if re.match(r"^(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z']*\s*\d{2,4}$", next_line, re.IGNORECASE):
                        continue
                    if not re.search(r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z\']*\s*\d{2,4}\s+to', next_line, re.IGNORECASE):
                        if any(kw in next_line for kw in ['Developer', 'Engineer', 'Analyst', 'Manager', 'Lead', 'Consultant', 'Architect', 'Admin', 'SSIS', 'MSBI', 'ETL', 'DBA', 'Database']):
                            title = next_line
                            title_line_idx = j
                            break
                        elif len(next_line) < 60 and next_line[0].isupper():
                            title = next_line
                            title_line_idx = j
                            break
                
                responsibilities = extract_responsibilities(lines, title_line_idx + 1, i + 50)
                
                duration = calculate_duration(start_year, start_month or 1, end_year, end_month or 12)
                start_date = f"{start_year}-{(start_month or 1):02d}"
                end_date = f"{end_year}-{(end_month or 12):02d}" if not is_present else f"{datetime.now().year}-{datetime.now().month:02d}"
                
                is_dup = any(
                    e.get('start_date') == start_date and
                    e.get('employer', '').lower() == employer.lower()
                    for e in experiences
                )
                
                if not is_dup:
                    experiences.append({
                        'employer': employer,
                        'title': title or None,
                        'start_date': start_date,
                        'end_date': end_date,
                        'duration_months': duration,
                        'responsibilities': responsibilities[:12],
                        'location': location,
                        'pattern': 8
                    })
        
        i += 1
    
    # =========================================================================
    # PATTERN 10: "## Title | Company, Location | Date" (Steven markdown)
    # =========================================================================
    steven_pattern = r'^\s*(?:##\s*)?(.+?)\s*\|\s*([^|,]+)(?:,\s*([^|]+))?\s*\|\s*((?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4})\s*[-–]+\s*((?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4}|Present|Current)'
    
    for i, line in enumerate(lines):
        match = re.match(steven_pattern, line.strip(), re.IGNORECASE)
        
        if match:
            title = match.group(1).strip()
            employer = match.group(2).strip()
            location = match.group(3).strip() if match.group(3) else None
            start_str = match.group(4)
            end_str = match.group(5)
            
            start_year, start_month, _ = parse_date(start_str)
            end_year, end_month, is_present = parse_date(end_str)
            
            if start_year and end_year:
                responsibilities = extract_responsibilities(lines, i + 1, i + 40)
                
                duration = calculate_duration(start_year, start_month or 1, end_year, end_month or 12)
                start_date = f"{start_year}-{(start_month or 1):02d}"
                end_date = f"{end_year}-{(end_month or 12):02d}" if not is_present else f"{datetime.now().year}-{datetime.now().month:02d}"
                
                is_dup = any(
                    e.get('start_date') == start_date and
                    e.get('employer', '').lower() == employer.lower()
                    for e in experiences
                )
                
                if not is_dup:
                    experiences.append({
                        'employer': employer,
                        'title': title,
                        'start_date': start_date,
                        'end_date': end_date,
                        'duration_months': duration,
                        'responsibilities': responsibilities[:12],
                        'location': location,
                        'pattern': 10
                    })
    
    # =========================================================================
    # PATTERN 7: "ROLE:" / "DESIGNATION:" keyword format (Sarwer)
    # =========================================================================
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if line and not line.startswith(('•', '-', '*', '■')):
            date_match = None
            for check_idx in range(max(0, i-1), min(len(lines), i+3)):
                check_line = lines[check_idx].strip()
                dm = re.search(date_pattern, check_line, re.IGNORECASE)
                if dm:
                    date_match = dm
                    break
            
            if date_match:
                for role_idx in range(i, min(len(lines), i + 8)):
                    role_line = lines[role_idx].strip()
                    role_match = re.match(r'^(?:ROLE|DESIGNATION)\s*[:\s]+\s*(.+)$', role_line, re.IGNORECASE)
                    
                    if role_match:
                        title = role_match.group(1).strip()
                        
                        employer = None
                        for emp_idx in range(role_idx - 1, max(0, role_idx - 6), -1):
                            emp_line = lines[emp_idx].strip()
                            if emp_line and not emp_line.startswith(('•', '-', '*', '■', 'ROLE', 'DESIGNATION')):
                                if 'work location' in emp_line.lower():
                                    continue
                                if re.match(r'^(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)', emp_line, re.IGNORECASE):
                                    continue
                                # Skip Client: lines - Pattern 12 handles those
                                if emp_line.lower().startswith('client:'):
                                    continue
                                if len(emp_line) > 3 and len(emp_line) < 100:
                                    if emp_line.upper() not in ['WORK EXPERIENCE', 'EXPERIENCE', 'PROFESSIONAL EXPERIENCE', 'ROLES & RESPONSIBILITIES']:
                                        dm = re.search(date_pattern, emp_line, re.IGNORECASE)
                                        if dm:
                                            employer = emp_line[:dm.start()].strip().rstrip('-–')
                                        else:
                                            employer = emp_line
                                        if employer and len(employer) > 2:
                                            break
                        
                        if employer and title:
                            start_str, end_str = date_match.group(1), date_match.group(2)
                            start_year, start_month, _ = parse_date(start_str)
                            end_year, end_month, is_present = parse_date(end_str)
                            
                            if start_year and end_year:
                                duration = calculate_duration(start_year, start_month or 1, end_year, end_month or 12)
                                start_date = f"{start_year}-{(start_month or 1):02d}"
                                end_date = f"{end_year}-{(end_month or 12):02d}" if not is_present else f"{datetime.now().year}-{datetime.now().month:02d}"
                                
                                responsibilities = extract_responsibilities(lines, role_idx + 1, role_idx + 20)
                                
                                is_dup = any(
                                    e.get('start_date') == start_date and
                                    (e.get('employer', '').lower() == employer.lower() or
                                     e.get('title', '').lower() == title.lower())
                                    for e in experiences
                                )
                                
                                if not is_dup:
                                    experiences.append({
                                        'employer': employer,
                                        'title': title,
                                        'start_date': start_date,
                                        'end_date': end_date,
                                        'duration_months': duration,
                                        'responsibilities': responsibilities[:12],
                                        'location': None,
                                        'pattern': 7
                                    })
                        break
        i += 1
    
    # =========================================================================
    # PATTERN 2: "Worked as X in Y from A to B" (Madhuri)
    # =========================================================================
    for i, line in enumerate(lines):
        if 'worked as' in line.lower():
            combined = line
            for j in range(i + 1, min(i + 4, len(lines))):
                next_line = lines[j].strip().lstrip('>')
                combined += ' ' + next_line
            combined = re.sub(r'\s+', ' ', combined)
            
            worked_pattern = r'[Ww]ork(?:ed|ing)\s+as\s+(?:a\s+)?(?:\*\*)?([^*]+?)(?:\*\*)?\s+(?:in|at)\s+(?:\*\*)?([^*]+?)(?:\*\*)?\s+from\s+([A-Za-z]+\s+\d{4})\s+to\s+([A-Za-z]+\s+\d{4}|Present)'
            
            m = re.search(worked_pattern, combined, re.IGNORECASE)
            if m:
                title = m.group(1).strip()
                employer = m.group(2).strip()
                start_str = m.group(3)
                end_str = m.group(4)
                
                start_year, start_month, _ = parse_date(start_str)
                end_year, end_month, is_present = parse_date(end_str)
                
                if start_year and end_year:
                    duration = calculate_duration(start_year, start_month or 1, end_year, end_month or 12)
                    start_date = f"{start_year}-{(start_month or 1):02d}"
                    end_date = f"{end_year}-{(end_month or 12):02d}" if not is_present else f"{datetime.now().year}-{datetime.now().month:02d}"
                    
                    responsibilities = extract_responsibilities(lines, i + 4, i + 30)
                    
                    is_dup = any(
                        e.get('start_date') == start_date and
                        (e.get('employer', '').lower() in employer.lower() or employer.lower() in e.get('employer', '').lower())
                        for e in experiences
                    )
                    
                    if not is_dup:
                        experiences.append({
                            'employer': employer,
                            'title': title,
                            'start_date': start_date,
                            'end_date': end_date,
                            'duration_months': duration,
                            'responsibilities': responsibilities,
                            'location': None,
                            'pattern': 2
                        })
    
    # =========================================================================
    # PATTERN 3: Table format "Client: X | Duration: Y" (Ramaswamy)
    # Handles multi-line tables where client/duration span multiple rows
    # =========================================================================
    i = 0
    while i < len(lines):
        line = lines[i]
        if 'client:' in line.lower() and 'duration' in ' '.join(lines[i:min(i+5, len(lines))]).lower():
            # Combine up to 10 lines to capture multi-line table cells
            combined = ' '.join(lines[i:min(i+10, len(lines))])
            combined = re.sub(r'[\*>|+=#]', ' ', combined)  # Remove table chars
            combined = re.sub(r'\s+', ' ', combined)  # Normalize whitespace
            
            # Extract client name - stop at Duration
            client_match = re.search(r'Client:\s*([A-Za-z][A-Za-z0-9\s&]+?)(?:\s+Duration|\s+Brief|\s*$)', combined, re.IGNORECASE)
            if not client_match:
                i += 1
                continue
            employer = client_match.group(1).strip()
            # Clean up employer name (remove trailing words like "Bank" that appear twice)
            employer = re.sub(r'\s+(Healthcare|Bank|Power|Water)\s*$', '', employer, flags=re.IGNORECASE).strip()
            
            # Extract duration - more flexible pattern
            # Allow extra words between "to" and end date
            duration_match = re.search(
                r'Duration[*:\s]*([A-Za-z]+)[-\s]*(\d{4})\s+to\s+(?:[A-Za-z]+\s+)?([A-Za-z]+[-\s]*\d{4}|Till\s*date|Present|Current)',
                combined, re.IGNORECASE
            )
            if not duration_match:
                # Try simpler pattern
                duration_match = re.search(
                    r'Duration[*:\s]*([A-Za-z]{3,9})[-\s]*(\d{4})\s+to\s+.*?([A-Za-z]{3,9})[-\s]*(\d{4})',
                    combined, re.IGNORECASE
                )
                if duration_match:
                    start_str = f"{duration_match.group(1)} {duration_match.group(2)}"
                    end_str = f"{duration_match.group(3)} {duration_match.group(4)}"
                else:
                    i += 1
                    continue
            else:
                start_str = f"{duration_match.group(1)} {duration_match.group(2)}"
                end_str = duration_match.group(3)
            
            start_year, start_month, _ = parse_date(start_str)
            end_year, end_month, is_present = parse_date(end_str)
            
            if start_year and end_year:
                duration = calculate_duration(start_year, start_month or 1, end_year, end_month or 12)
                
                # Try to extract title/role from nearby lines
                title = "Consultant"
                role_patterns = [
                    r'((?:Senior\s+)?(?:Data|Software|Cloud|GCP|ETL|BI|Snowflake)\s+(?:Engineer|Analyst|Developer|Architect|Consultant))',
                    r'((?:Lead|Sr\.?|Senior)\s+\w+\s+(?:Engineer|Developer|Analyst))',
                    r'Role[:\s]+([A-Za-z\s]+(?:Engineer|Developer|Analyst|Consultant))'
                ]
                for pattern in role_patterns:
                    role_match = re.search(pattern, combined, re.IGNORECASE)
                    if role_match:
                        title = role_match.group(1).strip()
                        break
                
                start_date = f"{start_year}-{(start_month or 1):02d}"
                end_date = f"{end_year}-{(end_month or 12):02d}" if not is_present else f"{datetime.now().year}-{datetime.now().month:02d}"
                
                # Extract responsibilities from table cells
                responsibilities = []
                for j in range(i + 1, min(i + 60, len(lines))):
                    resp_line = lines[j]
                    # Stop at next Client entry
                    if 'Client:' in resp_line and j > i + 2:
                        break
                    # Look for content in table cells
                    if '|' in resp_line:
                        cells = re.findall(r'\|\s*([^|]+)', resp_line)
                        for cell in cells:
                            cell = re.sub(r'^[>\s-]+', '', cell).strip()
                            if len(cell) > 30 and not cell.startswith(('+', '=', '-')):
                                if not re.match(r'^(Client|Duration|Brief|Role|Environment|Tech)', cell, re.IGNORECASE):
                                    responsibilities.append(cell[:500])
                
                is_dup = any(
                    e.get('start_date') == start_date and
                    (e.get('employer', '').lower() in employer.lower() or employer.lower() in e.get('employer', '').lower())
                    for e in experiences
                )
                
                if not is_dup:
                    experiences.append({
                        'employer': employer,
                        'title': title,
                        'start_date': start_date,
                        'end_date': end_date,
                        'duration_months': duration,
                        'responsibilities': responsibilities[:12],
                        'location': None,
                        'pattern': 3
                    })
            
            i += 5  # Skip ahead past the table we just processed
        else:
            i += 1
    
    # =========================================================================
    # PATTERN 1 & 5 & 6: Standard date-based extraction (fallback)
    # =========================================================================
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        line_stripped = re.sub(r'^[>\-•*#]+\s*', '', line).strip()
        
        if not line_stripped or line_stripped.startswith(('•', '-', '*', '■')):
            i += 1
            continue
        
        # Skip Client: format lines - Pattern 12 handles those better
        if line_stripped.lower().startswith('client:'):
            i += 1
            continue
        
        date_match = re.search(date_pattern, line_stripped, re.IGNORECASE)
        
        if date_match:
            if re.search(r'\bwork(?:ed|ing)\s+(?:as\s+)?(?:a\s+)?', line_stripped, re.IGNORECASE):
                i += 1
                continue
            
            edu_keywords = ['university', 'college', 'bachelor', 'master', 'mba', 'education', 'degree']
            if any(kw in line.lower() for kw in edu_keywords):
                i += 1
                continue
            
            header_preview = line_stripped[:date_match.start()].strip().lower()
            if header_preview in ['work experience', 'experience', 'employment', 'professional experience']:
                i += 1
                continue
            
            start_str, end_str = date_match.group(1), date_match.group(2)
            start_year, start_month, _ = parse_date(start_str)
            end_year, end_month, is_present = parse_date(end_str)
            
            if start_year and end_year:
                employer, title = "", ""
                header = line_stripped[:date_match.start()].strip()
                
                if '|' in header:
                    parts = [p.strip() for p in header.split('|') if p.strip()]
                    if len(parts) >= 1:
                        title = parts[0].strip()
                    
                    if not employer:
                        for back in range(1, min(10, i + 1)):
                            prev = lines[i - back].strip()
                            prev = re.sub(r'^[>\-•*#]+\s*', '', prev).strip()
                            if not prev or prev.startswith(('•', '-', '*')):
                                continue
                            if re.search(date_pattern, prev):
                                continue
                            if prev.upper() in ['PROFESSIONAL EXPERIENCE', 'WORK EXPERIENCE', 'EXPERIENCE']:
                                continue
                            if '–' in prev or ' - ' in prev:
                                employer = re.split(r'\s*[-–]+\s*', prev)[0].strip()
                            else:
                                employer = prev
                            break
                
                elif '–' in header or '--' in header or ' - ' in header:
                    parts = re.split(r'\s+[-–]+\s+|\s*--\s*', header, maxsplit=1)
                    if len(parts) == 2:
                        part1, part2 = parts[0].strip(), parts[1].strip()
                        
                        if any(kw in part1 for kw in title_keywords):
                            title, employer = part1, part2
                        elif any(kw in part2 for kw in title_keywords):
                            employer, title = part1, part2
                        else:
                            employer, title = part1, part2
                        
                        loc_pattern = rf'\s+({LOCATION_PATTERN})[,\s]*(PA|TX|IL|GA|NY|CA|OH|India|USA)?$'
                        if title:
                            title = re.sub(loc_pattern, '', title, flags=re.IGNORECASE).strip()
                        if employer:
                            employer = re.sub(loc_pattern, '', employer, flags=re.IGNORECASE).strip()
                    else:
                        employer = header
                
                else:
                    employer = header
                    
                    if i + 1 < len(lines):
                        next_line = lines[i + 1].strip()
                        next_clean = re.sub(r'^[>\-•*#]+\s*', '', next_line).strip()
                        
                        if next_clean and not next_clean.startswith(('•', '-', '*', '■')):
                            if not re.search(date_pattern, next_clean, re.IGNORECASE):
                                if any(kw.lower() in next_clean.lower() for kw in title_keywords):
                                    loc_pattern = rf'\s+({LOCATION_PATTERN})[,\s]*(PA|TX|IL|GA|NY|CA|OH|India|USA)?$'
                                    title = re.sub(loc_pattern, '', next_clean, flags=re.IGNORECASE).strip()
                
                if employer:
                    loc_pattern = rf'\s*[-–]+\s*({LOCATION_PATTERN})[,\s]*(PA|TX|IL|USA|India)?$'
                    employer = re.sub(loc_pattern, '', employer, flags=re.IGNORECASE).strip()
                    employer = re.sub(r'^Client[:\s]+', '', employer, flags=re.IGNORECASE).strip()
                
                responsibilities = extract_responsibilities(lines, i + 1, i + 30)
                
                duration = calculate_duration(start_year, start_month or 1, end_year, end_month or 12)
                start_date = f"{start_year}-{(start_month or 1):02d}"
                end_date = f"{end_year}-{(end_month or 12):02d}" if not is_present else f"{datetime.now().year}-{datetime.now().month:02d}"
                
                is_dup = any(
                    e.get('start_date') == start_date and
                    ((e.get('employer') or '').lower() == (employer or '').lower() or
                     (e.get('title') or '').lower() == (title or '').lower())
                    for e in experiences
                )
                
                if not is_dup and (employer or title):
                    experiences.append({
                        'employer': employer.strip() if employer else None,
                        'title': title.strip() if title else None,
                        'start_date': start_date,
                        'end_date': end_date,
                        'duration_months': duration,
                        'responsibilities': responsibilities[:12],
                        'location': None,
                        'pattern': 1
                    })
        i += 1
    
    # =========================================================================
    # PATTERN 11: "Project: X / Client: Y / Duration: Z" (Chakri format)
    # =========================================================================
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Look for "Project" line with tab or colon
        project_match = re.match(r'^Project\s*[:\t]+\s*(.+)$', line, re.IGNORECASE)
        if project_match:
            project_name = project_match.group(1).strip()
            
            # Look for Client and Duration in next few lines
            employer = None
            start_year = end_year = start_month = end_month = None
            is_present = False
            
            for j in range(i+1, min(i+5, len(lines))):
                next_line = lines[j].strip()
                
                # Client line
                client_match = re.match(r'^Client\s*[:\t]+\s*(.+)$', next_line, re.IGNORECASE)
                if client_match:
                    employer = client_match.group(1).strip()
                
                # Duration line
                duration_match = re.match(r'^Duration\s*[:\t]*\s*(.+)$', next_line, re.IGNORECASE)
                if duration_match:
                    date_str = duration_match.group(1).strip()
                    # Parse dates like "April 2021 to SEP 2024"
                    date_range = re.search(
                        r'(\w+)\s*(\d{4})\s*(?:to|[-–])\s*(\w+)\s*(\d{4}|Present|Current|Till)',
                        date_str, re.IGNORECASE
                    )
                    if date_range:
                        start_year, start_month, _ = parse_date(f"{date_range.group(1)} {date_range.group(2)}")
                        end_str = date_range.group(4)
                        if end_str.lower() in ['present', 'current', 'till']:
                            end_year = datetime.now().year
                            end_month = datetime.now().month
                            is_present = True
                        else:
                            end_year, end_month, is_present = parse_date(f"{date_range.group(3)} {date_range.group(4)}")
            
            # Look for responsibilities
            resp_start = i + 1
            for j in range(i+1, min(i+10, len(lines))):
                if re.match(r'^(Contribution|Responsibilities)', lines[j].strip(), re.IGNORECASE):
                    resp_start = j + 1
                    break
            
            responsibilities = extract_responsibilities(lines, resp_start, i + 25)
            
            if start_year and end_year:
                duration = calculate_duration(start_year, start_month or 1, end_year, end_month or 12)
                start_date = f"{start_year}-{(start_month or 1):02d}"
                end_date = f"{end_year}-{(end_month or 12):02d}" if not is_present else f"{datetime.now().year}-{datetime.now().month:02d}"
                
                is_dup = any(
                    e.get('start_date') == start_date and
                    (e.get('employer') or '').lower() == (employer or '').lower()
                    for e in experiences
                )
                
                if not is_dup and (employer or project_name):
                    experiences.append({
                        'employer': employer if employer else project_name,
                        'title': project_name if employer else None,
                        'start_date': start_date,
                        'end_date': end_date,
                        'duration_months': duration,
                        'responsibilities': responsibilities[:12],
                        'location': None,
                        'pattern': 11
                    })
        i += 1
    
    # =========================================================================
    # PATTERN 12: "Client: X / Role: Y" (Venkata/SAP consultant format)
    # Also handles "Client: X, Location \t\t Date Range" (Richie format)
    # =========================================================================
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Look for "Client:" line
        client_match = re.match(r'^Client:\s*(.+)$', line, re.IGNORECASE)
        if client_match:
            full_client_line = client_match.group(1).strip()
            employer = full_client_line
            title = None
            start_year = end_year = start_month = end_month = None
            is_present = False
            resp_start_line = i + 1
            
            # First check if dates are on the SAME line as Client:
            # Format: "Caterpillar Inc., Peoria, IL\t\t\tJuly 2021 – Oct 2023"
            same_line_date = re.search(
                r'(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s*(\d{4})\s*[-–to]+\s*(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?|Present|Current|Till)\w*\s*(\d{4})?',
                full_client_line, re.IGNORECASE
            )
            
            if same_line_date:
                # Extract employer (before the date)
                employer = full_client_line[:same_line_date.start()].strip().rstrip('\t')
                start_year, start_month, _ = parse_date(f"{same_line_date.group(1)} {same_line_date.group(2)}")
                end_str = same_line_date.group(3)
                if end_str.lower() in ['present', 'current', 'till']:
                    end_year = datetime.now().year
                    end_month = datetime.now().month
                    is_present = True
                else:
                    end_year_str = same_line_date.group(4) or str(datetime.now().year)
                    end_year, end_month, is_present = parse_date(f"{end_str} {end_year_str}")
            
            # Look for Role/Title and dates in next few lines
            for j in range(i+1, min(i+8, len(lines))):
                next_line = lines[j].strip()
                
                # Role line
                role_match = re.match(r'^(?:Role|Senior|Sr\.?|Jr\.?)\s*[:]?\s*(.+)$', next_line, re.IGNORECASE)
                if role_match and not title:
                    potential_title = role_match.group(1).strip() if next_line.startswith(('Role', 'role')) else next_line
                    # Verify it's a title, not a description
                    if len(potential_title) < 60 and not potential_title.startswith(('Project', 'Description', 'Responsibilities')):
                        title = potential_title
                        resp_start_line = j + 1
                
                # Title on line after Client (if no Role: prefix)
                if j == i + 1 and not title and not next_line.startswith(('Project', 'Description', 'Role', 'Responsibilities')):
                    if len(next_line) < 60 and re.search(r'(Consultant|Engineer|Developer|Manager|Analyst|Lead|Architect)', next_line, re.IGNORECASE):
                        title = next_line
                        resp_start_line = j + 1
                
                # Check for "Responsibilities:" header
                if re.match(r'^Responsibilities\s*:', next_line, re.IGNORECASE):
                    resp_start_line = j + 1
                
                # Date patterns (if not found on same line)
                if not start_year:
                    date_match = re.search(
                        r'(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s*(\d{4})\s*(?:to|[-–])\s*(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?|Present|Current|Till)\s*(\d{4})?',
                        next_line, re.IGNORECASE
                    )
                    if date_match:
                        start_year, start_month, _ = parse_date(f"{date_match.group(1)} {date_match.group(2)}")
                        end_str = date_match.group(3)
                        if end_str.lower() in ['present', 'current', 'till']:
                            end_year = datetime.now().year
                            end_month = datetime.now().month
                            is_present = True
                        else:
                            end_year_str = date_match.group(4) or str(datetime.now().year)
                            end_year, end_month, is_present = parse_date(f"{end_str} {end_year_str}")
            
            # If no dates found but this looks like current job (first Client: entry), assume Present
            if not start_year:
                # Check if this might be the most recent job (no dates = current)
                if i < 50:  # Near top of experience section
                    start_year = datetime.now().year - 1  # Assume started last year
                    start_month = 1
                    end_year = datetime.now().year
                    end_month = datetime.now().month
                    is_present = True
                else:
                    i += 1
                    continue
            
            # Extract responsibilities - start from after title/role line
            responsibilities = extract_responsibilities(lines, resp_start_line, resp_start_line + 30)
            
            if start_year and end_year:
                duration = calculate_duration(start_year, start_month or 1, end_year, end_month or 12)
                start_date = f"{start_year}-{(start_month or 1):02d}"
                end_date = f"{end_year}-{(end_month or 12):02d}" if not is_present else f"{datetime.now().year}-{datetime.now().month:02d}"
                
                # Check for duplicates
                is_dup = any(
                    e.get('start_date') == start_date and
                    (e.get('employer') or '').lower() in (employer or '').lower()
                    for e in experiences
                )
                
                if not is_dup and employer:
                    experiences.append({
                        'employer': employer,
                        'title': title,
                        'start_date': start_date,
                        'end_date': end_date,
                        'duration_months': duration,
                        'responsibilities': responsibilities[:12],
                        'location': None,
                        'pattern': 12
                    })
        i += 1
    
        # =========================================================================
    # PATTERN 13: Pipe-separated format "Date | Employer, Location | Title" (Jacques format)
    # =========================================================================
    # Example: "Jun. 2015 –Present | Fabrinet, Santa Clara, CA | Solution Architect"
    pipe_pattern = re.compile(
        r'^(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[.\s]*\s*(\d{4})\s*[–\-—to]+\s*'
        r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|Present|Current)[.\s]*\s*(\d{4})?\s*\|\s*'
        r'([^|]+)\s*\|\s*(.+)$',
        re.IGNORECASE
    )
    
    for i, line in enumerate(lines):
        match = pipe_pattern.match(line.strip())
        if match:
            start_month_str = match.group(1)
            start_year = int(match.group(2))
            end_month_str = match.group(3)
            end_year_str = match.group(4)
            employer_loc = match.group(5).strip()
            title = match.group(6).strip()
            
            # Parse employer and location
            if ',' in employer_loc:
                parts = employer_loc.rsplit(',', 1)
                employer = parts[0].strip()
                location = parts[-1].strip() if len(parts) > 1 else None
            else:
                employer = employer_loc
                location = None
            
            # Parse dates
            start_year, start_month, _ = parse_date(f"{start_month_str} {start_year}")
            
            if end_month_str.lower() in ['present', 'current']:
                end_year = datetime.now().year
                end_month = datetime.now().month
                is_present = True
            else:
                end_year = int(end_year_str) if end_year_str else datetime.now().year
                end_year, end_month, is_present = parse_date(f"{end_month_str} {end_year}")
            
            # Extract responsibilities from following lines
            responsibilities = extract_responsibilities(lines, i + 1, min(i + 50, len(lines)))
            
            duration = calculate_duration(start_year, start_month or 1, end_year, end_month or 12)
            start_date = f"{start_year}-{(start_month or 1):02d}"
            end_date = f"{end_year}-{(end_month or 12):02d}"
            
            # Check for duplicates
            is_dup = any(
                e.get('start_date') == start_date and
                employer.lower() in (e.get('employer') or '').lower()
                for e in experiences
            )
            
            if not is_dup:
                experiences.append({
                    'employer': employer,
                    'title': title,
                    'start_date': start_date,
                    'end_date': end_date,
                    'duration_months': duration,
                    'responsibilities': responsibilities[:12],
                    'location': location,
                    'pattern': 13
                })
    
    # =========================================================================
    # PATTERN 14: Tab-separated "Employer, Location \t\t Date" (Parth format)
    # =========================================================================
    # Example: "Accenture, New York, NY\t\t\t Jan 2019 - Present"
    tab_pattern = re.compile(
        r'^([A-Za-z][A-Za-z0-9\s&.,/-]+?)\s*[,]\s*([A-Za-z\s]+,?\s*[A-Z]{2})?\s*[\t]{2,}\s*'
        r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s*(\d{4})\s*[-–—to]+\s*'
        r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|Present|Current)[a-z]*\s*(\d{4})?',
        re.IGNORECASE
    )
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        match = tab_pattern.match(line)
        
        if match:
            employer = match.group(1).strip()
            location = match.group(2).strip() if match.group(2) else None
            start_month_str = match.group(3)
            start_year = int(match.group(4))
            end_month_str = match.group(5)
            end_year_str = match.group(6)
            
            # Get title from next line
            title = None
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                # Title line should not be a section header or responsibility
                if next_line and not next_line.startswith(('Key ', 'Role:', '•', '-', '*')):
                    if not re.match(r'^(EDUCATION|SKILLS|CERTIFICATION|SUMMARY)', next_line, re.IGNORECASE):
                        title = next_line
            
            # Parse dates
            start_year, start_month, _ = parse_date(f"{start_month_str} {start_year}")
            
            if end_month_str.lower() in ['present', 'current']:
                end_year = datetime.now().year
                end_month = datetime.now().month
                is_present = True
            else:
                end_year = int(end_year_str) if end_year_str else datetime.now().year
                end_year, end_month, is_present = parse_date(f"{end_month_str} {end_year}")
            
            # Extract responsibilities - look for "Key Responsibilities" section
            resp_start = i + 2
            responsibilities = extract_responsibilities(lines, resp_start, min(resp_start + 40, len(lines)))
            
            duration = calculate_duration(start_year, start_month or 1, end_year, end_month or 12)
            start_date = f"{start_year}-{(start_month or 1):02d}"
            end_date = f"{end_year}-{(end_month or 12):02d}"
            
            # Check for duplicates
            is_dup = any(
                e.get('start_date') == start_date and
                employer.lower() in (e.get('employer') or '').lower()
                for e in experiences
            )
            
            if not is_dup:
                experiences.append({
                    'employer': employer,
                    'title': title,
                    'start_date': start_date,
                    'end_date': end_date,
                    'duration_months': duration,
                    'responsibilities': responsibilities[:12],
                    'location': location,
                    'pattern': 14
                })
        i += 1
    
    
    
        # Sort by start date (most recent first)
    experiences.sort(key=lambda x: x.get('start_date', ''), reverse=True)
    
    # Remove pattern field and filter/clean invalid entries
    valid_experiences = []
    seen = set()
    
    for exp in experiences:
        exp.pop('pattern', None)
        exp.pop('client', None)
        
        employer = exp.get('employer') or ''
        
        # Skip entries with "Duration:" as employer
        if employer.lower().startswith('duration'):
            continue
        
        # Skip entries without both employer and title
        if not employer and not exp.get('title'):
            continue
        
        # Clean "Client:" prefix from employer names
        if employer.lower().startswith('client:'):
            employer = employer[7:].strip()
            exp['employer'] = employer
        
        # Remove duplicate entries (same employer base + similar date)
        # Normalize employer: take first word or before comma
        employer_norm = employer.split(',')[0].strip().lower()
        employer_first_word = employer_norm.split()[0] if employer_norm.split() else ''
        start_date = exp.get('start_date', '')
        
        # Check for duplicates using normalized employer
        is_dup = False
        for seen_emp, seen_date in seen:
            # Same start date and similar employer name
            if seen_date == start_date:
                if (seen_emp == employer_norm or 
                    seen_emp.startswith(employer_first_word) or 
                    employer_norm.startswith(seen_emp.split()[0] if seen_emp.split() else '')):
                    is_dup = True
                    break
        
        if is_dup:
            continue
        seen.add((employer_norm, start_date))
        
        # Normalize 'employer' to 'Employer' for output consistency
        if 'employer' in exp:
            exp['Employer'] = exp.pop('employer')
        
        valid_experiences.append(exp)
    
    return valid_experiences


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                    EXTRACTION AGENT - EDUCATION                              ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

def extract_education(text: str) -> List[Dict]:
    """Extract education from resume text (v8.6)."""
    education = []
    seen_degrees = set()
    
    # First check for pipe-separated format: "Education | Degree at/from University"
    for line in text.split('\n'):
        line = line.strip()
        if line.lower().startswith('education') and '|' in line:
            parts = line.split('|', 1)
            if len(parts) >= 2:
                edu_text = parts[1].strip()
                # Parse "Master of Computer Application (M.C.A) at Osmania University, India"
                at_match = re.search(r'(.+?)\s+at\s+(.+)', edu_text, re.IGNORECASE)
                from_match = re.search(r'(.+?)\s+from\s+(.+)', edu_text, re.IGNORECASE)
                
                if at_match:
                    education.append({
                        'degree': at_match.group(1).strip(),
                        'institution': at_match.group(2).strip(),
                        'year': None
                    })
                elif from_match:
                    education.append({
                        'degree': from_match.group(1).strip(),
                        'institution': from_match.group(2).strip(),
                        'year': None
                    })
                else:
                    education.append({
                        'degree': edu_text,
                        'institution': None,
                        'year': None
                    })
    
    if education:
        return education
    
    edu_match = re.search(
        r'(?:EDUCATION(?:AL)?\s*(?:QUALIFICATION|BACKGROUND)?|ACADEMIC\s*(?:QUALIFICATION)?)'
        r'[:\s]*\n?(.+?)(?:\nPROFESSIONAL|\nWORK|\nTECHNICAL|\nSKILLS|\nCERTIFI|\nEXPERIENCE|\nCLIENT:|\Z)',
        text, re.IGNORECASE | re.DOTALL
    )
    
    if edu_match:
        edu_section = edu_match.group(1)
        lines = [l.strip() for l in edu_section.split('\n') if l.strip()]
        
        for line in lines:
            # Strip bullet points
            line = re.sub(r'^[•·\-\*]\s*', '', line).strip()
            if not line:
                continue
            # Skip non-education lines
            if re.match(r'^(Worked|Working|PROFESSIONAL|Client:|Duration:|Tools|Brief|Role|Environment)', line, re.IGNORECASE):
                continue
            # Skip lines that look like job descriptions
            if 'Bank' in line and 'Duration' in line:
                continue
            if re.match(r'^\|', line):
                continue
            
            entry = {}
            
            if line.count('|') >= 2:
                parts = [p.strip() for p in line.split('|')]
                entry['degree'] = parts[0]
                # Find year and institution - year is typically in middle, institution at end
                entry['institution'] = parts[-1] if not re.match(r'^\d{4}$', parts[-1].strip()) else None
                for p in parts[1:]:
                    year_match = re.search(r'(\d{4})', p)
                    if year_match:
                        entry['year'] = year_match.group(1)
                        # If institution not set and this isn't the year-only part
                        if not entry.get('institution') and len(parts) >= 3:
                            entry['institution'] = parts[-1]
                        break
            
            elif '|' in line and ('–' in line or '-' in line):
                dash_match = re.match(r'^(.+?)\s*[-–]\s*(.+?)\s*\|\s*(\d{4})\s*[-–]\s*(\d{4})', line)
                if dash_match:
                    entry['degree'] = dash_match.group(1).strip()
                    entry['institution'] = dash_match.group(2).strip()
                    entry['year'] = dash_match.group(4)
            
            elif re.search(r'\bfrom\b', line, re.IGNORECASE):
                from_match = re.match(r'^(.+?)\s+[Ff]rom\s+(.+?)\s+(\d{4})', line)
                if from_match:
                    entry['degree'] = from_match.group(1).strip()
                    entry['institution'] = from_match.group(2).strip()
                    entry['year'] = from_match.group(3)
            
            else:
                degree_patterns = [r'\bmaster', r'\bbachelor', r'\bmba\b', r'\bmca\b', 
                                  r'\bb\.?tech\b', r'\bm\.?tech\b', r'\bb\.?e\b', r'\bm\.?e\b']
                if any(re.search(p, line.lower()) for p in degree_patterns):
                    parts = re.split(r'\s*[-–]\s*', line, maxsplit=1)
                    entry['degree'] = parts[0].strip()
                    if len(parts) > 1:
                        year_match = re.search(r'(\d{4})\s*$', parts[1])
                        if year_match:
                            entry['year'] = year_match.group(1)
                            entry['institution'] = parts[1][:year_match.start()].strip().rstrip(',|')
                        else:
                            entry['institution'] = parts[1].strip()
            
            if entry.get('degree'):
                degree_key = entry['degree'].lower()[:50]
                if degree_key not in seen_degrees:
                    seen_degrees.add(degree_key)
                    education.append({
                        'degree': entry.get('degree'),
                        'institution': entry.get('institution'),
                        'year': entry.get('year')
                    })
    
    # Post-process: Extract institution from degree string if missing
    for edu in education:
        if not edu.get('institution') and edu.get('degree'):
            degree = edu['degree']
            
            # Pattern 1: "..., University of X, ..."
            univ_match = re.search(r',\s*(University\s+of\s+[^,]+|[A-Z][a-z]+\s+University[^,]*)', degree)
            if univ_match:
                edu['institution'] = univ_match.group(1).strip().rstrip('.')
                edu['degree'] = degree[:univ_match.start()].strip().rstrip(',')
                continue
            
            # Pattern 2: "... from University"
            from_match = re.search(r'(.+?)\s+(?:from|at)\s+([^,]+University[^,]*)', degree, re.IGNORECASE)
            if from_match:
                edu['degree'] = from_match.group(1).strip()
                edu['institution'] = from_match.group(2).strip().rstrip('.')
                continue
            
            # Pattern 3: Just look for University in comma-separated parts
            parts = [p.strip() for p in degree.split(',')]
            for i, part in enumerate(parts):
                if 'university' in part.lower() or 'college' in part.lower() or 'institute' in part.lower():
                    edu['institution'] = part.strip().rstrip('.')
                    edu['degree'] = ', '.join(parts[:i]).strip()
                    break
    
    return education[:10]


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                  EXTRACTION AGENT - CERTIFICATIONS                           ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

def extract_certifications(text: str) -> List[str]:
    """Extract certifications from resume text (v8.6)."""
    certifications = []
    
    # Check for certification patterns in first 40 lines
    for line in text.split('\n')[:40]:
        line = line.strip()
        
        # Pattern 1: "Certified Consultant/Professional/Expert..."
        if re.match(r'^Certified\s+(Consultant|Professional|Expert|Scrum|AWS|Azure|Google|SAP|Business|Data|Cloud|Solution|Associate|Developer|Administrator)', line, re.IGNORECASE):
            cert = re.sub(r'^[•·\-\*]\s*', '', line).strip()
            if cert and len(cert) > 10 and cert not in certifications:
                certifications.append(cert)
        
        # Pattern 2: "AWS/Azure/Google/SAP Certified..."
        elif re.match(r'^(AWS|Azure|Google|Google Cloud|SAP|Microsoft|Oracle|Cisco|PMP|ITIL|Scrum)\s+Certified', line, re.IGNORECASE):
            cert = re.sub(r'^[•·\-\*]\s*', '', line).strip()
            if cert and len(cert) > 10 and cert not in certifications:
                certifications.append(cert)
        
        # Pattern 3: Contains "Certified" and common cert keywords
        elif 'certified' in line.lower() and any(kw in line.lower() for kw in ['cloud', 'engineer', 'architect', 'developer', 'administrator', 'associate', 'professional', 'scrum', 'pmp', 'aws', 'azure', 'gcp']):
            cert = re.sub(r'^[•·\-\*]\s*', '', line).strip()
            # Clean up trailing periods
            cert = cert.rstrip('.')
            if cert and len(cert) > 10 and cert not in certifications:
                certifications.append(cert)
    
    # Then check for CERTIFICATIONS section
    cert_match = re.search(
        r'CERTIFICATIONS?(?:\s*/\s*TRAININGS?)?[:\s]*\n(.+?)(?:\nPROFESSIONAL|\nEXPERIENCE|\nEDUCATION|\nSKILLS|\nTOOLS|\Z)',
        text, re.IGNORECASE | re.DOTALL
    )
    
    if cert_match:
        for line in cert_match.group(1).split('\n'):
            line = re.sub(r'^[•·\-\*]\s*', '', line.strip())
            if line and 3 < len(line) < 200:
                if re.match(r'^(PROFESSIONAL|EXPERIENCE|IMG|IBM|Cognizant)', line, re.IGNORECASE):
                    continue
                line = re.sub(r'\s*[-–]\s*In Progress\.?$', '', line, flags=re.IGNORECASE)
                if line not in certifications:
                    certifications.append(line)
    
    return list(dict.fromkeys(certifications))[:20]


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                     EXTRACTION AGENT - SKILLS (v8.0)                         ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

def extract_skills(text: str) -> Dict[str, List[str]]:
    """Extract and categorize skills from resume text."""
    found_skills = {cat: [] for cat in SKILL_CATEGORIES}
    text_lower = text.lower()
    
    for category, skills in SKILL_CATEGORIES.items():
        for skill in skills:
            if re.search(rf'\b{re.escape(skill)}\b', text_lower):
                if len(skill) <= 4 and skill.lower() not in ['go', 'r']:
                    display = skill.upper()
                elif skill in ['node.js', 'asp.net', 'vb.net', '.net']:
                    display = skill
                else:
                    display = skill.title()
                if display not in found_skills[category]:
                    found_skills[category].append(display)
    
    return {k: v for k, v in found_skills.items() if v}


def calculate_skill_experience(text: str, experiences: List[Dict]) -> Dict[str, List[Dict]]:
    """
    Calculate skill experience months based on job mentions.
    v8.0: Added role-based skill inference.
    """
    skill_months: Dict[str, int] = {}
    found_skills: Dict[str, str] = {}
    text_lower = text.lower()
    
    # Find all skills mentioned in the resume
    for category, skills in SKILL_CATEGORIES.items():
        for skill in skills:
            if re.search(rf'\b{re.escape(skill)}\b', text_lower):
                display = skill.upper() if len(skill) <= 4 and skill.lower() not in ['go', 'r'] else skill.title()
                found_skills[skill] = display
                skill_months[skill] = 0
    
    # Calculate experience for each skill
    for exp in experiences:
        exp_text = ' '.join([
            str(exp.get('title', '')),
            str(exp.get('employer', '') or ''),
            ' '.join(exp.get('responsibilities', []))
        ]).lower()
        
        # Direct skill mention
        for skill in found_skills:
            if re.search(rf'\b{re.escape(skill)}\b', exp_text):
                skill_months[skill] += exp.get('duration_months', 0)
        
        # v8.0: Role-based skill inference
        title_lower = str(exp.get('title', '')).lower()
        for role_key, role_skills in ROLE_SKILL_MAP.items():
            if role_key in title_lower:
                for skill in role_skills:
                    if skill in found_skills and skill_months[skill] == 0:
                        skill_months[skill] = exp.get('duration_months', 0)
    
    # Build result
    result = {cat: [] for cat in SKILL_CATEGORIES}
    for category, skills in SKILL_CATEGORIES.items():
        for skill in skills:
            if skill in found_skills:
                result[category].append({
                    'skill': found_skills[skill],
                    'experience_months': skill_months.get(skill, 0)
                })
        result[category].sort(key=lambda x: x['experience_months'], reverse=True)
    
    return {k: v for k, v in result.items() if v}


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                EXTRACTION AGENT - TITLE & SUMMARY                            ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

def extract_title(text: str, experiences: List[Dict]) -> str:
    """Extract professional title from summary or experience."""
    summary_match = re.search(
        r'(?:PROFESSIONAL\s+)?SUMMARY[:\s]*\n(.+?)(?:\n[A-Z]{2,}|\Z)',
        text, re.IGNORECASE | re.DOTALL
    )
    
    if summary_match:
        summary = summary_match.group(1)
        
        title_match = re.match(
            r'^([\w\s/]+(?:Manager|Engineer|Developer|Analyst|Consultant|Architect|Lead))\s+with\s+\d+',
            summary.strip(), re.IGNORECASE
        )
        if title_match:
            return title_match.group(1).strip()
        
        title_match = re.search(
            r'\d+\+?\s+years?\s+(?:of\s+)?experience\s+(?:as\s+(?:a|an)\s+|in\s+)([\w\s]+?)(?:\.|,|and)',
            summary, re.IGNORECASE
        )
        if title_match:
            title = title_match.group(1).strip()
            if any(kw in title.lower() for kw in ['manager', 'engineer', 'developer', 'analyst']):
                return title
    
    if experiences and experiences[0].get('title'):
        return experiences[0]['title']
    
    return ""


def extract_summary(text: str) -> str:
    """Extract professional summary."""
    patterns = [
        r'(?:PROFESSIONAL\s+)?SUMMARY[:\s]*\n(.+?)(?:\nSKILLS|\nEXPERIENCE|\nWORK|\nTECHNICAL|\Z)',
        r'EXPERIENCE\s+SUMMARY[:\s]*\n(.+?)(?:\nTECHNICAL|\nSKILLS|\nPROFESSIONAL|\Z)',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            summary = match.group(1).strip()
            summary = re.sub(r'\n+', ' ', summary)
            return summary[:2000]
    
    return ""


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║              VALIDATION AGENT (v8.5 - Comprehensive Rules)                   ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

"""
VALIDATION RULES:
═══════════════════════════════════════════════════════════════════════════════

CRITICAL FIELDS (must have):
  1. name         - Full name, 2+ chars, not a job title/skill
  2. email        - Valid email format with @
  3. experience   - At least 1 job entry

IMPORTANT FIELDS (should have):
  4. phone        - Phone number
  5. title        - Current/recent job title
  6. firstname    - First name extracted
  7. lastname     - Last name extracted

EXPERIENCE QUALITY:
  8. Each job should have: Employer, title, start_date, end_date
  9. Each job should have: 3+ responsibilities
  10. Dates should be valid (start < end)

EDUCATION:
  11. At least 1 education entry (degree + institution)
  12. Year should be numeric 4 digits

SKILLS:
  13. At least 5 technical skills extracted

CERTIFICATIONS:
  14. Extract if present in resume

LOCATION:
  15. Extract city/state if present

AI ENHANCEMENT TRIGGERS:
- Any CRITICAL field missing → AI
- Score < 70 → AI
- Low responsibilities (< 5 total) → AI
- Missing education when text mentions degree → AI
- Missing certifications when text mentions "certified" → AI
"""

def validation_agent(parsed: Dict, text: str) -> ValidationResult:
    """
    Comprehensive validation with detailed rules (v8.6.1).
    
    VALIDATION RULES:
    ═══════════════════════════════════════════════════════════════════════════
    CRITICAL (15 pts each):
      1. name         - Must exist, 2+ chars, not a job title
      2. email        - Valid format with @
      3. experience   - At least 1 job
    
    IMPORTANT (8 pts each):
      4. phone        - Phone number
      5. title        - Job title
    
    IDENTITY (3 pts each):
      6. firstname    - First name extracted
      7. lastname     - Last name extracted
    
    EXPERIENCE QUALITY:
      8. responsibilities  - 5+ total (10 pts), 10+ expected (5 pts)
      9. employers         - More than half should have employer
      10. job_titles       - More than half should have title
      11. dates            - More than half should have dates
    
    CONTENT:
      12. education        - When text has degree keywords (5 pts)
      13. certifications   - When text has "certified" (3 pts)
      14. location         - When text has location patterns (2 pts)
      15. skills           - At least 5 technical skills (3 pts)
      16. skill_experience - Should have non-zero values (2 pts)
    
    QUALITY:
      17. duplicate_employers    - Same employer multiple times (3 pts)
      18. excessive_education    - More than 5 entries likely error (5 pts)
      19. single_word_name       - Only firstname, no lastname (3 pts)
    
    AI ENHANCEMENT TRIGGERS:
      - Any critical field missing
      - Score below 70
      - quality_issues: low_responsibilities, missing_employers, duplicate_employers
      - content_gaps: missing_education, missing_certifications, single_word_name
    """
    result = ValidationResult()
    pr = parsed.get("parsed_resume", {})
    text_lower = text.lower()
    
    # ═══════════════════════════════════════════════════════════════════════════
    # CRITICAL CHECKS (15 points each)
    # ═══════════════════════════════════════════════════════════════════════════
    
    # 1. Name validation
    name = pr.get("name", "") or ""
    if not name or len(name) < 3:
        result.issues.append("missing_name")
        result.missing_fields.append("name")
        result.score -= 15
    else:
        # Check for invalid names (job titles, skills, etc.)
        invalid_names = [
            "sql server", "data engineer", "project manager", "key expertise", 
            "professional", "technical", "summary", "experience", "education",
            "developer", "analyst", "consultant", "manager", "engineer",
            "skills", "certification", "objective", "resume", "curriculum"
        ]
        name_lower = name.lower()
        if any(inv in name_lower for inv in invalid_names) and len(name.split()) < 2:
            result.issues.append("invalid_name")
            result.missing_fields.append("name")
            result.score -= 15
    
    # 2. Email validation
    email = pr.get("email", "") or ""
    if not email or "@" not in email or "." not in email.split("@")[-1]:
        result.issues.append("missing_email")
        result.missing_fields.append("email")
        result.score -= 15
    
    # 3. Experience validation
    experiences = pr.get("experience", [])
    if not experiences:
        result.issues.append("missing_experience")
        result.missing_fields.append("experience")
        result.score -= 15
    
    # ═══════════════════════════════════════════════════════════════════════════
    # IMPORTANT CHECKS (8 points each)
    # ═══════════════════════════════════════════════════════════════════════════
    
    # 4. Phone validation
    if not pr.get("phone_number"):
        result.issues.append("missing_phone")
        result.missing_fields.append("phone")
        result.score -= 8
    
    # 5. Title validation
    if not pr.get("title"):
        result.issues.append("missing_title")
        result.missing_fields.append("title")
        result.score -= 8
    
    # 6-7. First/Last name validation
    if not pr.get("firstname"):
        result.issues.append("missing_firstname")
        result.missing_fields.append("firstname")
        result.score -= 3
    
    if not pr.get("lastname"):
        result.issues.append("missing_lastname")
        result.missing_fields.append("lastname")
        result.score -= 3
    
    # ═══════════════════════════════════════════════════════════════════════════
    # EXPERIENCE QUALITY (up to 20 points)
    # ═══════════════════════════════════════════════════════════════════════════
    
    if experiences:
        total_resp = sum(len(e.get("responsibilities", [])) for e in experiences)
        jobs_without_employer = sum(1 for e in experiences if not (e.get("employer") or e.get("Employer")))
        jobs_without_title = sum(1 for e in experiences if not e.get("title"))
        jobs_without_dates = sum(1 for e in experiences if not e.get("start_date"))
        jobs_without_resp = sum(1 for e in experiences if not e.get("responsibilities"))
        
        # 9. Responsibilities check
        if total_resp < 5:
            result.issues.append("low_responsibilities")
            result.missing_fields.append("responsibilities")
            result.score -= 10
        elif total_resp < 10:
            result.issues.append("few_responsibilities")
            result.score -= 5
        
        # 8. Job details check
        if jobs_without_employer > len(experiences) // 2:
            result.issues.append("missing_employers")
            result.score -= 5
        
        if jobs_without_title > len(experiences) // 2:
            result.issues.append("missing_job_titles")
            result.score -= 5
        
        if jobs_without_dates > len(experiences) // 2:
            result.issues.append("missing_dates")
            result.score -= 3
    
    # ═══════════════════════════════════════════════════════════════════════════
    # EDUCATION CHECK (5 points)
    # ═══════════════════════════════════════════════════════════════════════════
    
    education = pr.get("education", [])
    education_keywords = ["bachelor", "master", "mba", "phd", "degree", "university", "college", "b.tech", "m.tech"]
    text_has_education = any(kw in text_lower for kw in education_keywords)
    
    if not education and text_has_education:
        result.issues.append("missing_education")
        result.missing_fields.append("education")
        result.score -= 5
    elif education:
        # Check education quality
        for edu in education:
            if not edu.get("institution"):
                result.issues.append("education_missing_institution")
                result.missing_fields.append("education")
                result.score -= 2
                break
    
    # ═══════════════════════════════════════════════════════════════════════════
    # CERTIFICATIONS CHECK (3 points)
    # ═══════════════════════════════════════════════════════════════════════════
    
    certifications = pr.get("certifications", [])
    cert_keywords = ["certified", "certification", "certificate", "aws certified", "pmp", "scrum master", "cissp"]
    text_has_certs = any(kw in text_lower for kw in cert_keywords)
    
    if not certifications and text_has_certs:
        result.issues.append("missing_certifications")
        result.missing_fields.append("certifications")
        result.score -= 3
    
    # ═══════════════════════════════════════════════════════════════════════════
    # LOCATION CHECK (2 points)
    # ═══════════════════════════════════════════════════════════════════════════
    
    location = pr.get("location")
    location_patterns = ["tx", "ca", "ny", "fl", "usa", "india", "uk", "houston", "dallas", "new york", "bangalore"]
    text_has_location = any(loc in text_lower for loc in location_patterns)
    
    if not location and text_has_location:
        result.issues.append("missing_location")
        result.missing_fields.append("location")
        result.score -= 2
    
    # ═══════════════════════════════════════════════════════════════════════════
    # SKILLS CHECK (3 points)
    # ═══════════════════════════════════════════════════════════════════════════
    
    skills = pr.get("technical_skills", [])
    if len(skills) < 5:
        result.issues.append("few_skills")
        result.score -= 3
    
    # ═══════════════════════════════════════════════════════════════════════════
    # SKILL EXPERIENCE CHECK
    # ═══════════════════════════════════════════════════════════════════════════
    
    key_skills = pr.get("key_skills", {})
    if key_skills:
        # Check if all skill experience is 0
        all_skills_zero = True
        for category_skills in key_skills.values():
            if isinstance(category_skills, list):
                for skill_info in category_skills:
                    if isinstance(skill_info, dict) and skill_info.get("experience_months", 0) > 0:
                        all_skills_zero = False
                        break
        
        if all_skills_zero and experiences:
            result.issues.append("skill_experience_not_calculated")
            result.score -= 2
    
    # ═══════════════════════════════════════════════════════════════════════════
    # DUPLICATE EMPLOYER CHECK
    # ═══════════════════════════════════════════════════════════════════════════
    
    if experiences:
        employer_counts = {}
        for exp in experiences:
            emp = (exp.get("Employer") or "").lower().strip()
            if emp:
                employer_counts[emp] = employer_counts.get(emp, 0) + 1
        
        duplicate_employers = [e for e, count in employer_counts.items() if count > 1]
        if duplicate_employers:
            result.issues.append("duplicate_employers")
            result.score -= 3
    
    # ═══════════════════════════════════════════════════════════════════════════
    # EXCESSIVE EDUCATION CHECK (likely parsing error)
    # ═══════════════════════════════════════════════════════════════════════════
    
    if len(education) > 5:
        result.issues.append("excessive_education_entries")
        result.score -= 5
    
    # ═══════════════════════════════════════════════════════════════════════════
    # SINGLE WORD NAME CHECK
    # ═══════════════════════════════════════════════════════════════════════════
    
    if name and len(name.split()) == 1:
        result.issues.append("single_word_name")
        result.missing_fields.append("lastname")
        result.score -= 3
    
    # ═══════════════════════════════════════════════════════════════════════════
    # AI ENHANCEMENT DECISION
    # ═══════════════════════════════════════════════════════════════════════════
    
    # Trigger AI if:
    # - Any critical field missing (name, email, experience)
    # - Score below 70
    # - Low responsibilities
    # - Missing education/certs when text suggests they exist
    # - Single-word name (missing lastname)
    # - Skill experience not calculated
    # - Duplicate employers detected
    # - Excessive education entries
    
    critical_missing = any(issue in result.issues for issue in [
        "missing_name", "invalid_name", "missing_email", "missing_experience"
    ])
    
    quality_issues = any(issue in result.issues for issue in [
        "low_responsibilities", "missing_employers", "missing_job_titles",
        "skill_experience_not_calculated", "duplicate_employers", "excessive_education_entries"
    ])
    
    content_gaps = any(issue in result.issues for issue in [
        "missing_education", "missing_certifications", "missing_location",
        "missing_lastname", "missing_firstname", "single_word_name"
    ])
    
    result.needs_ai_enhancement = (
        critical_missing or 
        quality_issues or 
        content_gaps or 
        result.score < 70
    )
    
    result.score = max(0, result.score)
    
    return result


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                    AI ENHANCEMENT AGENT (v8.5 - Comprehensive)               ║
# ╚══════════════════════════════════════════════════════════════════════════════╝


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                    AGENTIC SELF-HEALING SYSTEM (v8.7.0)                      ║
# ╠══════════════════════════════════════════════════════════════════════════════╣
# ║  This system ensures the parser works with ANY resume format by:             ║
# ║  1. Pattern Detection - Identifies resume format patterns                    ║
# ║  2. Gap Analysis - Compares detected vs extracted fields                     ║
# ║  3. Targeted AI Fix - Uses AI only for missing fields                        ║
# ║  4. Quality Validation - Ensures AI output meets standards                   ║
# ║  5. Failure Logging - Records issues for future pattern improvement          ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

class AgenticOrchestrator:
    """
    Central orchestrator for the agentic resume parsing system.
    Coordinates between detection, extraction, diagnosis, and AI enhancement.
    """
    
    def __init__(self, text: str, filename: str):
        self.text = text
        self.filename = filename
        self.diagnosis = None
        self.extraction_result = None
        self.ai_attempts = 0
        self.max_ai_attempts = 2
        self.failure_log = []
    
    def detect_format(self) -> Dict:
        """
        Agent 1: Format Detection
        Analyzes the resume to understand its structure.
        """
        format_info = {
            "type": "unknown",
            "patterns_detected": [],
            "structure": {},
            "confidence": 0.0
        }
        
        lines = self.text.split('\n')[:50]  # First 50 lines
        
        # Detect pipe-separated format (Jacques style)
        pipe_lines = sum(1 for l in lines if '|' in l and re.search(r'\d{4}', l))
        if pipe_lines >= 3:
            format_info["patterns_detected"].append("pipe_separated")
            format_info["type"] = "pipe_structured"
        
        # Detect Client:/Role: format
        client_lines = sum(1 for l in lines if re.match(r'^Client:', l, re.I))
        if client_lines >= 2:
            format_info["patterns_detected"].append("client_role")
            format_info["type"] = "consultant_format"
        
        # Detect Project:/Duration: format
        project_lines = sum(1 for l in lines if re.match(r'^Project\s*:', l, re.I))
        if project_lines >= 2:
            format_info["patterns_detected"].append("project_duration")
            format_info["type"] = "project_based"
        
        # Detect table format (tabs/multiple spaces)
        tab_lines = sum(1 for l in lines if '\t' in l or '    ' in l)
        if tab_lines >= 5:
            format_info["patterns_detected"].append("tabular")
        
        # Detect standard chronological
        date_lines = sum(1 for l in lines if re.search(r'(19|20)\d{2}\s*[-–to]+\s*(19|20|Present|Current)', l, re.I))
        if date_lines >= 3:
            format_info["patterns_detected"].append("chronological")
        
        # Calculate confidence
        format_info["confidence"] = min(1.0, len(format_info["patterns_detected"]) * 0.3)
        
        return format_info
    
    def analyze_gaps(self) -> Dict:
        """
        Agent 2: Gap Analysis
        Compares what patterns detected vs what was actually extracted.
        """
        if not self.diagnosis:
            self.diagnosis = diagnose_extraction(self.text, self.filename)
        
        gaps = {
            "critical_missing": [],
            "important_missing": [],
            "minor_missing": [],
            "extraction_failures": []
        }
        
        results = self.diagnosis.get("extraction_results", {})
        patterns = self.diagnosis.get("raw_patterns_found", {})
        
        # Critical fields
        if not results.get("name") or len(str(results.get("name", "")).split()) < 2:
            if patterns.get("name", {}).get("all_caps_name") or patterns.get("name", {}).get("title_case_name"):
                gaps["extraction_failures"].append("name")
            gaps["critical_missing"].append("name")
        
        if not results.get("email"):
            if patterns.get("contact", {}).get("email"):
                gaps["extraction_failures"].append("email")
            gaps["critical_missing"].append("email")
        
        if results.get("experience_count", 0) == 0:
            if any(patterns.get("experience", {}).values()):
                gaps["extraction_failures"].append("experience")
            gaps["critical_missing"].append("experience")
        
        # Important fields
        if not results.get("phone"):
            if patterns.get("contact", {}).get("phone_standard") or patterns.get("contact", {}).get("phone_intl"):
                gaps["extraction_failures"].append("phone")
            gaps["important_missing"].append("phone")
        
        if results.get("education_count", 0) == 0:
            if patterns.get("education", {}).get("degree_keyword"):
                gaps["extraction_failures"].append("education")
            gaps["important_missing"].append("education")
        
        if results.get("total_responsibilities", 0) < 10:
            gaps["important_missing"].append("responsibilities")
        
        # Minor fields
        if results.get("certification_count", 0) == 0:
            if patterns.get("certifications", {}).get("certified_keyword"):
                gaps["extraction_failures"].append("certifications")
            gaps["minor_missing"].append("certifications")
        
        return gaps
    
    def build_targeted_prompt(self, gaps: Dict) -> str:
        """
        Agent 3: Prompt Builder
        Creates a focused prompt for AI based on specific gaps.
        """
        prompt_parts = []
        
        # Add context about what we need
        prompt_parts.append("Extract ONLY the following missing information from this resume:")
        prompt_parts.append("")
        
        # Critical missing
        if "name" in gaps.get("critical_missing", []):
            prompt_parts.append("1. FULL NAME (first and last name, look at the very top of the resume)")
        
        if "experience" in gaps.get("critical_missing", []) or "experience" in gaps.get("extraction_failures", []):
            prompt_parts.append("""
2. WORK EXPERIENCE - Extract ALL jobs with:
   - employer: Company name
   - title: Job title
   - start_date: YYYY-MM format
   - end_date: YYYY-MM format (or "Present")
   - responsibilities: List of 3-5 key responsibilities
   
   Look for patterns like:
   - "Company Name | Title | Date"
   - "Client: Company" followed by "Role: Title"
   - Company names followed by dates
""")
        
        # Important missing
        if "phone" in gaps.get("important_missing", []):
            prompt_parts.append("3. PHONE NUMBER (any format: +1, parentheses, dashes)")
        
        if "education" in gaps.get("important_missing", []):
            prompt_parts.append("""
4. EDUCATION - Extract:
   - degree: Full degree name
   - institution: University/College name
   - year: Graduation year
""")
        
        if "responsibilities" in gaps.get("important_missing", []):
            prompt_parts.append("""
5. JOB RESPONSIBILITIES - For each job, extract:
   - Key achievements and duties
   - Technical skills used
   - Projects completed
   Look for bullet points, action verbs (Led, Developed, Implemented, etc.)
""")
        
        prompt_parts.append("")
        prompt_parts.append("Return ONLY a JSON object with the extracted fields. Do not include fields that are already extracted or not found.")
        
        return "\n".join(prompt_parts)
    
    async def run_targeted_ai_fix(self, gaps: Dict, current_result: Dict) -> Dict:
        """
        Agent 4: Targeted AI Fixer
        Uses AI to extract only missing fields.
        """
        if not ANTHROPIC_API_KEY:
            return current_result
        
        if self.ai_attempts >= self.max_ai_attempts:
            self.failure_log.append(f"Max AI attempts ({self.max_ai_attempts}) reached")
            return current_result
        
        self.ai_attempts += 1
        
        # Build targeted prompt
        prompt = self.build_targeted_prompt(gaps)
        
        # Prepare text excerpt (first 4000 chars for context)
        text_excerpt = self.text[:4000]
        
        try:
            async with httpx.AsyncClient(timeout=60.0) as client:
                response = await client.post(
                    "https://api.anthropic.com/v1/messages",
                    headers={
                        "x-api-key": ANTHROPIC_API_KEY,
                        "anthropic-version": "2023-06-01",
                        "content-type": "application/json"
                    },
                    json={
                        "model": "claude-sonnet-4-20250514",
                        "max_tokens": 2000,
                        "messages": [{
                            "role": "user",
                            "content": f"{prompt}\n\n---RESUME TEXT---\n{text_excerpt}"
                        }]
                    }
                )
                
                if response.status_code == 200:
                    data = response.json()
                    ai_text = data.get("content", [{}])[0].get("text", "")
                    
                    # Parse AI response
                    ai_result = self._parse_ai_response(ai_text)
                    
                    # Merge with current result
                    current_result = self._merge_ai_result(current_result, ai_result, gaps)
                    
        except Exception as e:
            self.failure_log.append(f"AI call failed: {str(e)}")
        
        return current_result
    
    def _parse_ai_response(self, ai_text: str) -> Dict:
        """Parse AI response, handling various formats."""
        try:
            # Try to find JSON in response
            json_match = re.search(r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}', ai_text, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
        except:
            pass
        
        # Fallback: extract key-value pairs
        result = {}
        
        # Extract name
        name_match = re.search(r'name["\'\s:]+([A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)', ai_text)
        if name_match:
            result["name"] = name_match.group(1)
        
        # Extract phone
        phone_match = re.search(r'phone["\'\s:]+([+\d\s\-().]+)', ai_text)
        if phone_match:
            result["phone"] = phone_match.group(1).strip()
        
        return result
    
    def _merge_ai_result(self, current: Dict, ai_result: Dict, gaps: Dict) -> Dict:
        """Merge AI result into current result, only for missing fields."""
        pr = current.get("parsed_resume", {})
        
        # Only update truly missing fields
        if "name" in gaps.get("critical_missing", []) and ai_result.get("name"):
            name_parts = ai_result["name"].split()
            if len(name_parts) >= 2:
                pr["firstname"] = name_parts[0]
                pr["lastname"] = name_parts[-1]
                pr["name"] = ai_result["name"]
        
        if "phone" in gaps.get("important_missing", []) and ai_result.get("phone"):
            pr["phone_number"] = ai_result["phone"]
        
        if ai_result.get("experience") and isinstance(ai_result["experience"], list):
            # Add AI-extracted experience only if we have very few
            if len(pr.get("experience", [])) < 3:
                for exp in ai_result["experience"]:
                    if isinstance(exp, dict) and exp.get("employer"):
                        pr["experience"].append({
                            "Employer": exp.get("employer"),
                            "title": exp.get("title"),
                            "start_date": exp.get("start_date"),
                            "end_date": exp.get("end_date"),
                            "responsibilities": exp.get("responsibilities", []),
                            "ai_extracted": True
                        })
        
        if ai_result.get("education") and isinstance(ai_result["education"], list):
            if len(pr.get("education", [])) == 0:
                pr["education"] = ai_result["education"]
        
        current["parsed_resume"] = pr
        current["ai_enhanced"] = True
        
        return current
    
    def log_failure(self, issue: str, context: Dict = None):
        """Log parsing failures for future improvement."""
        self.failure_log.append({
            "issue": issue,
            "filename": self.filename,
            "context": context,
            "timestamp": datetime.now().isoformat()
        })
    
    def get_failure_report(self) -> List:
        """Get all logged failures."""
        return self.failure_log




# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                    QUALITY ASSURANCE SYSTEM                                  ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

class QualityAssurance:
    """
    Ensures parsing quality meets standards.
    Implements automatic retry and escalation logic.
    """
    
    MIN_ACCEPTABLE_SCORE = 70
    MIN_GOOD_SCORE = 85
    
    CRITICAL_FIELDS = ["name", "email", "experience"]
    IMPORTANT_FIELDS = ["phone_number", "education", "technical_skills"]
    
    @classmethod
    def assess_quality(cls, result: Dict) -> Dict:
        """Assess the quality of parsing result."""
        pr = result.get("parsed_resume", {})
        score = result.get("validation_score", 0)
        
        assessment = {
            "score": score,
            "grade": cls._score_to_grade(score),
            "critical_fields_present": {},
            "important_fields_present": {},
            "recommendations": [],
            "needs_retry": False,
            "needs_manual_review": False
        }
        
        # Check critical fields
        for field in cls.CRITICAL_FIELDS:
            value = pr.get(field)
            if field == "experience":
                present = len(value or []) > 0
            else:
                present = bool(value)
            assessment["critical_fields_present"][field] = present
            
            if not present:
                assessment["recommendations"].append(f"Missing critical field: {field}")
                assessment["needs_retry"] = True
        
        # Check important fields
        for field in cls.IMPORTANT_FIELDS:
            value = pr.get(field)
            if isinstance(value, list):
                present = len(value) > 0
            else:
                present = bool(value)
            assessment["important_fields_present"][field] = present
        
        # Check experience quality
        experiences = pr.get("experience", [])
        if experiences:
            total_resp = sum(len(e.get("responsibilities", [])) for e in experiences)
            if total_resp < 5:
                assessment["recommendations"].append("Very few responsibilities extracted")
                assessment["needs_retry"] = True
            
            employers_missing = sum(1 for e in experiences if not (e.get("Employer") or e.get("employer")))
            if employers_missing > len(experiences) // 2:
                assessment["recommendations"].append("Many jobs missing employer names")
                assessment["needs_retry"] = True
        
        # Determine if manual review needed
        if score < 50:
            assessment["needs_manual_review"] = True
            assessment["recommendations"].append("Score too low - may need manual review or new pattern")
        
        return assessment
    
    @classmethod
    def _score_to_grade(cls, score: int) -> str:
        if score >= 95:
            return "A+"
        elif score >= 90:
            return "A"
        elif score >= 85:
            return "B+"
        elif score >= 80:
            return "B"
        elif score >= 70:
            return "C"
        elif score >= 60:
            return "D"
        else:
            return "F"
    
    @classmethod
    async def ensure_quality(cls, text: str, filename: str, max_retries: int = 2) -> Dict:
        """
        Main quality assurance function.
        Retries with AI if quality is insufficient.
        """
        # First attempt without AI
        result = await parse_resume(text, filename, use_ai=False)
        assessment = cls.assess_quality(result)
        
        retry_count = 0
        
        while assessment["needs_retry"] and retry_count < max_retries and ANTHROPIC_API_KEY:
            retry_count += 1
            
            # Use agentic parsing with AI
            result = await agentic_parse(text, filename, use_ai=True)
            assessment = cls.assess_quality(result)
            
            # If score improved significantly, we're done
            if result.get("validation_score", 0) >= cls.MIN_ACCEPTABLE_SCORE:
                break
        
        result["quality_assessment"] = assessment
        result["retry_count"] = retry_count
        
        return result


async def smart_parse(text: str, filename: str) -> Dict:
    """
    Smart parsing with automatic quality assurance.
    This is the recommended function for production use.
    """
    return await QualityAssurance.ensure_quality(text, filename)



async def agentic_parse(text: str, filename: str, use_ai: bool = True) -> Dict:
    """
    Main agentic parsing function.
    Orchestrates all agents to ensure robust parsing.
    """
    orchestrator = AgenticOrchestrator(text, filename)
    
    # Step 1: Detect format
    format_info = orchestrator.detect_format()
    
    # Step 2: Run standard extraction
    result = await parse_resume(text, filename, use_ai=False)
    
    # Step 3: Analyze gaps
    gaps = orchestrator.analyze_gaps()
    
    # Step 4: If gaps exist and AI is enabled, run targeted fix
    has_critical_gaps = len(gaps.get("critical_missing", [])) > 0
    has_extraction_failures = len(gaps.get("extraction_failures", [])) > 0
    
    if use_ai and (has_critical_gaps or has_extraction_failures):
        result = await orchestrator.run_targeted_ai_fix(gaps, result)
        
        # Re-validate after AI fix
        validation = validation_agent(result, text)
        result["validation_score"] = validation.score
        result["validation_issues"] = validation.issues
    
    # Add metadata
    result["agentic_metadata"] = {
        "format_detected": format_info,
        "gaps_analyzed": gaps,
        "ai_attempts": orchestrator.ai_attempts,
        "failures": orchestrator.failure_log
    }
    
    return result



async def ai_enhancement_agent(text: str, parsed: Dict, validation: ValidationResult) -> Dict:
    """
    Use Claude API to extract/fix ALL missing or low-quality fields.
    Comprehensive extraction for any validation failures (v8.5).
    """
    if not ANTHROPIC_API_KEY:
        parsed["ai_skipped"] = "No API key configured"
        return parsed
    
    try:
        import httpx
        
        pr = parsed.get("parsed_resume", {})
        missing = list(set(validation.missing_fields))
        issues = validation.issues
        
        # Build comprehensive prompt based on what's needed
        existing_jobs = pr.get("experience", [])
        
        # Run diagnosis to understand what patterns exist
        try:
            diagnosis = diagnose_extraction(text, "")
            raw_patterns = diagnosis.get("raw_patterns_found", {})
        except:
            raw_patterns = {}
        job_list = ""
        if existing_jobs:
            job_list = "EXISTING JOBS FOUND:\n" + "\n".join([
                f"- {j.get('employer', 'Unknown')}: {j.get('title', 'Unknown')} ({j.get('start_date', '?')} - {j.get('end_date', '?')})" 
                for j in existing_jobs
            ])
        
        # Determine what needs extraction
        needs_name = "name" in missing or "firstname" in missing or "lastname" in missing
        needs_contact = "email" in missing or "phone" in missing or "location" in missing
        needs_responsibilities = "responsibilities" in missing or "low_responsibilities" in issues
        needs_education = "education" in missing
        needs_certifications = "certifications" in missing
        needs_experience = "experience" in missing or "missing_employers" in issues
        
        prompt = f"""You are an expert resume parser. Extract the following MISSING information from this resume.
Return ONLY valid JSON with no markdown formatting or explanation.

FIELDS TO EXTRACT: {', '.join(missing) if missing else 'Validate and enhance all fields'}

ISSUES DETECTED: {', '.join(issues) if issues else 'None'}

{job_list}

RESUME TEXT:
{text[:15000]}

Return this exact JSON structure (include ALL fields, use null if not found):
{{
  "firstname": "First name",
  "lastname": "Last name",
  "name": "Full Name",
  "email": "email@example.com",
  "phone": "phone number",
  "location": "City, State",
  "title": "Current/Most Recent Job Title",
  "certifications": ["cert1", "cert2"],
  "education": [
    {{"degree": "Degree Name", "institution": "University/College Name", "year": "YYYY"}}
  ],
  "experience": [
    {{
      "employer": "Company Name",
      "title": "Job Title",
      "start_date": "YYYY-MM",
      "end_date": "YYYY-MM or Present",
      "location": "City, State",
      "responsibilities": ["responsibility 1", "responsibility 2", "..."]
    }}
  ]
}}

IMPORTANT:
- Extract the candidate's ACTUAL name from the resume, not job titles
- For responsibilities, extract 5-15 bullet points per job
- For education, include institution name and graduation year
- For certifications, include full certification names
- Match employers exactly to existing jobs when adding responsibilities"""

        async with httpx.AsyncClient(timeout=120.0) as client:
            response = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": ANTHROPIC_API_KEY,
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json"
                },
                json={
                    "model": "claude-sonnet-4-5-20250929",
                    "max_tokens": 4096,
                    "messages": [{"role": "user", "content": prompt}]
                }
            )
            
            if response.status_code == 200:
                ai_text = response.json().get("content", [{}])[0].get("text", "")
                json_match = re.search(r'\{[\s\S]*\}', ai_text)
                
                if json_match:
                    ai_data = json.loads(json_match.group())
                    fixed_fields = []
                    
                    # ═══════════════════════════════════════════════════════════
                    # MERGE AI DATA INTO PARSED RESULT
                    # ═══════════════════════════════════════════════════════════
                    
                    # Name fields
                    if ai_data.get("name") and not pr.get("name"):
                        pr["name"] = ai_data["name"]
                        fixed_fields.append("name")
                    
                    if ai_data.get("firstname") and not pr.get("firstname"):
                        pr["firstname"] = ai_data["firstname"]
                        fixed_fields.append("firstname")
                    elif ai_data.get("name") and not pr.get("firstname"):
                        parts = ai_data["name"].split()
                        if parts:
                            pr["firstname"] = parts[0]
                            fixed_fields.append("firstname")
                    
                    if ai_data.get("lastname") and not pr.get("lastname"):
                        pr["lastname"] = ai_data["lastname"]
                        fixed_fields.append("lastname")
                    elif ai_data.get("name") and not pr.get("lastname"):
                        parts = ai_data["name"].split()
                        if len(parts) > 1:
                            pr["lastname"] = parts[-1]
                            fixed_fields.append("lastname")
                    
                    # Contact fields
                    if ai_data.get("email") and not pr.get("email"):
                        pr["email"] = ai_data["email"]
                        fixed_fields.append("email")
                    
                    if ai_data.get("phone") and not pr.get("phone_number"):
                        pr["phone_number"] = ai_data["phone"]
                        fixed_fields.append("phone")
                    
                    if ai_data.get("location") and not pr.get("location"):
                        pr["location"] = ai_data["location"]
                        fixed_fields.append("location")
                    
                    if ai_data.get("title") and not pr.get("title"):
                        pr["title"] = ai_data["title"]
                        fixed_fields.append("title")
                    
                    # Education
                    if ai_data.get("education"):
                        existing_edu = pr.get("education", [])
                        ai_edu = ai_data["education"]
                        if isinstance(ai_edu, list):
                            if not existing_edu or len(ai_edu) > len(existing_edu):
                                pr["education"] = ai_edu
                                fixed_fields.append("education")
                            else:
                                # Merge missing fields
                                for i, edu in enumerate(existing_edu):
                                    if i < len(ai_edu):
                                        if not edu.get("institution") and ai_edu[i].get("institution"):
                                            edu["institution"] = ai_edu[i]["institution"]
                                            fixed_fields.append("education_institution")
                                        if not edu.get("year") and ai_edu[i].get("year"):
                                            edu["year"] = ai_edu[i]["year"]
                                            fixed_fields.append("education_year")
                    
                    # Certifications
                    if ai_data.get("certifications"):
                        existing_certs = pr.get("certifications", [])
                        ai_certs = ai_data["certifications"]
                        if isinstance(ai_certs, list) and len(ai_certs) > len(existing_certs):
                            pr["certifications"] = ai_certs
                            fixed_fields.append("certifications")
                    
                    # Experience - merge responsibilities
                    if ai_data.get("experience") and existing_jobs:
                        ai_jobs = ai_data["experience"]
                        for existing_job in existing_jobs:
                            existing_employer = (existing_job.get("Employer") or "").lower()
                            existing_resp = existing_job.get("responsibilities", [])
                            
                            # Find matching AI job
                            for ai_job in ai_jobs:
                                ai_employer = (ai_job.get("employer") or "").lower()
                                
                                # Fuzzy match on employer
                                if existing_employer and ai_employer:
                                    match = (
                                        existing_employer in ai_employer or 
                                        ai_employer in existing_employer or
                                        any(word in ai_employer for word in existing_employer.split() if len(word) > 3)
                                    )
                                    
                                    if match:
                                        ai_resp = ai_job.get("responsibilities", [])
                                        # Add responsibilities if we have fewer
                                        if ai_resp and len(ai_resp) > len(existing_resp):
                                            existing_job["responsibilities"] = ai_resp
                                            if "responsibilities" not in fixed_fields:
                                                fixed_fields.append("responsibilities")
                                        
                                        # Fill missing employer
                                        if not existing_job.get("Employer") and ai_job.get("employer"):
                                            existing_job["Employer"] = ai_job["employer"]
                                            fixed_fields.append("employer")
                                        
                                        # Fill missing location
                                        if not existing_job.get("location") and ai_job.get("location"):
                                            existing_job["location"] = ai_job["location"]
                                        
                                        break
                    
                    # If no existing jobs, use AI jobs entirely
                    elif ai_data.get("experience") and not existing_jobs:
                        ai_jobs = ai_data["experience"]
                        pr["experience"] = [
                            {
                                "Employer": j.get("employer"),
                                "title": j.get("title"),
                                "start_date": j.get("start_date"),
                                "end_date": j.get("end_date"),
                                "location": j.get("location"),
                                "responsibilities": j.get("responsibilities", []),
                                "duration_months": 0
                            }
                            for j in ai_jobs
                        ]
                        fixed_fields.append("experience")
                    
                    parsed["ai_enhanced"] = True
                    parsed["ai_fields_fixed"] = list(set(fixed_fields))
            else:
                parsed["ai_error"] = f"API {response.status_code}: {response.text[:200]}"
                    
    except Exception as e:
        parsed["ai_error"] = str(e)
    
    return parsed


# ║                           OUTPUT AGENT                                       ║
# ╚══════════════════════════════════════════════════════════════════════════════╝



# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                    SELF-DIAGNOSIS & DEBUG SYSTEM (v8.6.3)                    ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

def diagnose_extraction(text: str, filename: str) -> Dict:
    """
    Comprehensive diagnosis of what was extracted and what's missing.
    Returns detailed report for debugging new formats.
    """
    diagnosis = {
        "filename": filename,
        "text_length": len(text),
        "text_preview": text[:500] if text else "",
        "extraction_results": {},
        "issues_found": [],
        "suggestions": [],
        "raw_patterns_found": {}
    }
    
    # Check what patterns exist in raw text
    text_lower = text.lower()
    
    # Name patterns
    name_patterns = {
        "name_email_line": bool(re.search(r'^[A-Z][a-z]+.*Email', text[:500], re.MULTILINE)),
        "all_caps_name": bool(re.search(r'^[A-Z]{2,}\s+[A-Z]{2,}', text[:300], re.MULTILINE)),
        "title_case_name": bool(re.search(r'^[A-Z][a-z]+\s+[A-Z][a-z]+', text[:300], re.MULTILINE)),
    }
    diagnosis["raw_patterns_found"]["name"] = name_patterns
    
    # Contact patterns
    contact_patterns = {
        "email": bool(re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)),
        "phone_standard": bool(re.search(r'\d{3}[-.]?\d{3}[-.]?\d{4}', text)),
        "phone_intl": bool(re.search(r'\+1.*\d{3}.*\d{3}.*\d{4}', text)),
        "linkedin": bool(re.search(r'linkedin', text_lower)),
    }
    diagnosis["raw_patterns_found"]["contact"] = contact_patterns
    
    # Experience patterns
    exp_patterns = {
        "client_role": bool(re.search(r'Client:\s*.+', text)),
        "project_client": bool(re.search(r'Project\s*[:\t]', text)),
        "company_dates": bool(re.search(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{4}\s*[-–to]', text, re.IGNORECASE)),
        "pipe_format": bool(re.search(r'[A-Za-z]+\s*\|\s*[A-Za-z]+\s*\|', text)),
        "date_range": bool(re.search(r'\d{4}\s*[-–]\s*(\d{4}|Present|Current)', text, re.IGNORECASE)),
    }
    diagnosis["raw_patterns_found"]["experience"] = exp_patterns
    
    # Education patterns  
    edu_patterns = {
        "education_section": bool(re.search(r'EDUCATION', text, re.IGNORECASE)),
        "degree_keyword": bool(re.search(r'(Bachelor|Master|MBA|MCA|B\.?Tech|M\.?Tech|Ph\.?D)', text, re.IGNORECASE)),
        "university": bool(re.search(r'University|College|Institute', text, re.IGNORECASE)),
        "pipe_education": bool(re.search(r'Education\s*\|', text, re.IGNORECASE)),
    }
    diagnosis["raw_patterns_found"]["education"] = edu_patterns
    
    # Certification patterns
    cert_patterns = {
        "certified_keyword": bool(re.search(r'Certified', text, re.IGNORECASE)),
        "certification_section": bool(re.search(r'CERTIFICATION', text, re.IGNORECASE)),
        "aws_cert": bool(re.search(r'AWS\s+Certified', text, re.IGNORECASE)),
        "pmp": bool(re.search(r'\bPMP\b', text)),
    }
    diagnosis["raw_patterns_found"]["certifications"] = cert_patterns
    
    # Now do actual extraction and compare
    contact = extract_contact(text)
    firstname, middle, lastname = extract_name(text, filename)
    experiences = extract_experiences(text)
    education = extract_education(text)
    certifications = extract_certifications(text)
    skills = extract_skills(text)
    
    diagnosis["extraction_results"] = {
        "name": f"{firstname} {middle} {lastname}".strip(),
        "email": contact.get("email"),
        "phone": contact.get("phone"),
        "experience_count": len(experiences),
        "education_count": len(education),
        "certification_count": len(certifications),
        "skill_count": sum(len(v) for v in skills.values()),
        "total_responsibilities": sum(len(e.get("responsibilities", [])) for e in experiences),
    }
    
    # Generate suggestions based on gaps
    if not firstname and name_patterns.get("all_caps_name"):
        diagnosis["suggestions"].append("Name appears to be in ALL CAPS format - check extract_name patterns")
    if not firstname and name_patterns.get("name_email_line"):
        diagnosis["suggestions"].append("Name + Email on same line detected - check Strategy 0a")
    
    if not contact.get("phone") and contact_patterns.get("phone_intl"):
        diagnosis["suggestions"].append("International phone format detected but not extracted - check phone_patterns")
    
    if len(experiences) == 0 and exp_patterns.get("client_role"):
        diagnosis["suggestions"].append("Client:/Role: format detected but no jobs extracted - check Pattern 12")
    if len(experiences) == 0 and exp_patterns.get("project_client"):
        diagnosis["suggestions"].append("Project:/Client: format detected but no jobs extracted - check Pattern 11")
    
    if len(education) == 0 and edu_patterns.get("degree_keyword"):
        diagnosis["suggestions"].append("Degree keywords found but no education extracted - check extract_education")
    if len(education) == 0 and edu_patterns.get("pipe_education"):
        diagnosis["suggestions"].append("Pipe-format education detected - check pipe-separated extraction")
    
    if len(certifications) == 0 and cert_patterns.get("certified_keyword"):
        diagnosis["suggestions"].append("'Certified' keyword found but not extracted - check extract_certifications")
    
    return diagnosis


def format_diagnosis_report(diagnosis: Dict) -> str:
    """Format diagnosis as readable report."""
    report = []
    report.append("=" * 80)
    report.append(f"DIAGNOSIS REPORT: {diagnosis['filename']}")
    report.append("=" * 80)
    
    report.append(f"\nText Length: {diagnosis['text_length']} chars")
    report.append(f"\nText Preview (first 300 chars):")
    report.append(diagnosis['text_preview'][:300])
    
    report.append("\n" + "-" * 40)
    report.append("RAW PATTERNS DETECTED:")
    report.append("-" * 40)
    for category, patterns in diagnosis['raw_patterns_found'].items():
        report.append(f"\n{category.upper()}:")
        for pattern, found in patterns.items():
            status = "✅" if found else "❌"
            report.append(f"  {status} {pattern}")
    
    report.append("\n" + "-" * 40)
    report.append("EXTRACTION RESULTS:")
    report.append("-" * 40)
    for field, value in diagnosis['extraction_results'].items():
        status = "✅" if value else "❌"
        report.append(f"  {status} {field}: {value}")
    
    if diagnosis['suggestions']:
        report.append("\n" + "-" * 40)
        report.append("SUGGESTIONS TO FIX:")
        report.append("-" * 40)
        for i, suggestion in enumerate(diagnosis['suggestions'], 1):
            report.append(f"  {i}. {suggestion}")
    
    return "\n".join(report)



async def parse_resume(text: str, filename: str = None, use_ai: bool = True) -> Dict:
    """
    Orchestrate all agents and produce final standardized output.
    Flow: Extraction → Validation → AI Enhancement → JSON Output
    """
    text = normalize_text(text)
    
    # ===== EXTRACTION AGENT =====
    contact = extract_contact(text)
    firstname, middle, lastname = extract_name(text, filename or "")
    experiences = extract_experiences(text)
    education = extract_education(text)
    certifications = extract_certifications(text)
    skills = extract_skills(text)
    skill_experience = calculate_skill_experience(text, experiences)
    summary = extract_summary(text)
    
    # Build full name
    name_parts = [firstname]
    if middle:
        name_parts.append(middle)
    name_parts.append(lastname)
    name = ' '.join(filter(None, name_parts))
    
    # Extract title
    title = extract_title(text, experiences)
    
    # Calculate totals
    total_months = sum(e.get('duration_months', 0) for e in experiences)
    
    # Build result
    result = {
        "parsed_resume": {
            "firstname": firstname or None,
            "lastname": lastname or None,
            "name": name or None,
            "title": title or None,
            "location": contact.get('location') or None,
            "phone_number": contact.get('phone') or None,
            "email": contact.get('email') or None,
            "linkedin": contact.get('linkedin') or None,
            "summary": summary or None,
            "total_experience_months": total_months,
            "total_experience_years": round(total_months / 12, 1) if total_months else 0,
            "technical_skills": [s for cat in skills.values() for s in cat],
            "key_skills": skill_experience,
            "education": education,
            "certifications": certifications,
            "experience": [
                {
                    "Employer": e.get('Employer') or e.get('employer'),
                    "title": e.get('title'),
                    "location": e.get('location'),
                    "start_date": e.get('start_date'),
                    "end_date": e.get('end_date'),
                    "duration_months": e.get('duration_months'),
                    "responsibilities": e.get('responsibilities', [])
                }
                for e in experiences
            ],
            "filename": filename or ""
        },
        "parser_version": VERSION
    }
    
    # ===== VALIDATION AGENT =====
    validation = validation_agent(result, text)
    result["validation_score"] = validation.score
    result["validation_issues"] = validation.issues
    result["needs_ai_enhancement"] = validation.needs_ai_enhancement
    
    # ===== AI ENHANCEMENT AGENT =====
    if use_ai and validation.needs_ai_enhancement and ANTHROPIC_API_KEY:
        result = await ai_enhancement_agent(text, result, validation)
    
    return result


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                           API ENDPOINTS                                      ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

@app.get("/", tags=["Info"])
async def root():
    """API information and capabilities."""
    return {
        "name": "Resume Parser API",
        "version": VERSION,
        "docs": "/docs",
        "redoc": "/redoc",
        "ai_validation": "enabled" if ANTHROPIC_API_KEY else "disabled",
        "supported_formats": ["pdf", "docx", "txt", "zip"],
        "architecture": "Agentic Framework v8.0",
        "improvements": [
            "Enhanced name extraction with filename fallback",
            "Role-based skill inference",
            "Short date parsing (Jul24, Feb'20)",
            "10 experience extraction patterns",
            "Two-column PDF handling",
            "Lenient validation scoring"
        ]
    }


@app.get("/health", tags=["Health"])
async def health_check():
    """Health check endpoint."""
    return {
        "status": "healthy",
        "ai_available": bool(ANTHROPIC_API_KEY),
        "version": VERSION
    }


@app.post("/parse/file", tags=["Parsing"])
async def parse_file(
    file: UploadFile = File(..., description="Resume file (PDF, DOCX, TXT)"),
    use_ai_validation: bool = Form(default=True, description="Enable AI enhancement")
):
    """Parse resume from uploaded file."""
    if not file.filename:
        raise HTTPException(400, "No file provided")
    
    try:
        content = await file.read()
        
        if len(content) == 0:
            raise HTTPException(400, "Empty file")
        
        text = extract_text_intelligent(content, file.filename)
        
        if len(text.strip()) < 50:
            file_type = detect_file_type(content, file.filename)
            raise HTTPException(
                400, 
                f"Insufficient text extracted ({len(text)} chars). "
                f"Detected file type: {file_type}"
            )
        
        result = await parse_resume(text, file.filename, use_ai_validation)
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Error processing file: {str(e)}")


@app.post("/parse", tags=["Parsing"])
@app.post("/parse/text", tags=["Parsing"])
async def parse_text(request: ParseTextRequest):
    """Parse resume from text input."""
    try:
        result = await parse_resume(
            request.text,
            request.filename,
            request.use_ai_validation
        )
        return result
    except Exception as e:
        raise HTTPException(500, f"Error: {str(e)}")


@app.post("/extract/skills", tags=["Utilities"])
async def extract_skills_endpoint(request: ExtractSkillsRequest):
    """Extract and categorize technical skills from text."""
    skills = extract_skills(request.text)
    flat_skills = [s for cat in skills.values() for s in cat]
    return {
        "skills": flat_skills, 
        "categorized": skills, 
        "count": len(flat_skills),
        "categories": list(skills.keys())
    }


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                    COMPATIBILITY EXPORTS FOR api_server.py                   ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

class ResponseFormat(str, Enum):
    """Response format enum for compatibility."""
    JSON = "json"
    MARKDOWN = "markdown"


class ParseResumeInput(BaseModel):
    """Input model for parse_resume_full - compatibility with api_server.py."""
    resume_text: str = Field(..., min_length=50, max_length=500000)
    response_format: ResponseFormat = Field(default=ResponseFormat.JSON)
    filename: Optional[str] = Field(default=None)
    file_path: Optional[str] = Field(default=None)
    use_ai_validation: bool = Field(default=True)


async def parse_resume_full(params: ParseResumeInput) -> str:
    """Compatibility wrapper for api_server.py imports."""
    result = await parse_resume(
        text=params.resume_text,
        filename=params.filename,
        use_ai=params.use_ai_validation
    )
    return json.dumps(result, indent=2, ensure_ascii=False)


def extract_technical_skills(text: str) -> List[str]:
    """Compatibility alias for extract_skills()."""
    categorized = extract_skills(text)
    return [skill for skills_list in categorized.values() for skill in skills_list]


# ╔══════════════════════════════════════════════════════════════════════════════╗
# ║                           MAIN ENTRY POINT                                   ║
# ╚══════════════════════════════════════════════════════════════════════════════╝

if __name__ == "__main__":
    import uvicorn
    
    port = int(os.environ.get("PORT", 8080))
    
    print("╔" + "═" * 70 + "╗")
    print(f"║{'RESUME PARSER API v' + VERSION + ' - Enterprise Agentic Framework':^70}║")
    print("╠" + "═" * 70 + "╣")
    print(f"║  Port: {port:<62}║")
    print(f"║  AI Enhancement: {'Enabled' if ANTHROPIC_API_KEY else 'Disabled':<52}║")
    print(f"║  Docs: http://localhost:{port}/docs{' ' * 37}║")
    print("╠" + "═" * 70 + "╣")
    print("║  v8.0 Improvements:                                                  ║")
    print("║    • Enhanced name extraction with filename fallback                 ║")
    print("║    • Role-based skill inference (Java Dev = Java experience)         ║")
    print("║    • Short date parsing (Jul24, Feb'20, May'14)                      ║")
    print("║    • 10 experience extraction patterns                               ║")
    print("║    • Two-column PDF handling                                         ║")
    print("║    • Lenient validation scoring                                      ║")
    print("╚" + "═" * 70 + "╝")
    
    uvicorn.run(app, host="0.0.0.0", port=port)


# DEBUG ENDPOINT
@app.post("/debug/text", tags=["Debug"])
async def debug_text_extraction(file: UploadFile = File(...)):
    """Debug endpoint - shows extracted text and name detection."""
    content = await file.read()
    text = extract_text_intelligent(content, file.filename)
    lines = text.split('\n')
    
    all_caps_lines = []
    for i, line in enumerate(lines):
        clean = ' '.join(line.split()).strip('\r')
        if clean.isupper() and 5 < len(clean) < 40:
            parts = clean.split()
            if 2 <= len(parts) <= 4:
                all_caps_lines.append({"line_num": i, "text": clean, "parts": parts})
    
    first, middle, last = extract_name(text, file.filename)
    
    return {
        "filename": file.filename,
        "text_length": len(text),
        "line_count": len(lines),
        "first_20_lines": [l.strip()[:80] for l in lines[:20]],
        "all_caps_candidates": all_caps_lines[:10],
        "extracted_name": {"first": first, "middle": middle, "last": last, "full": ' '.join(filter(None, [first, middle, last]))}
    }@app.post("/parse/smart")
async def parse_resume_smart(
    file: UploadFile = File(...)
):
    """
    Smart resume parsing with automatic quality assurance.
    
    This endpoint:
    1. Tries pattern-based extraction first
    2. Automatically retries with AI if quality is low
    3. Returns quality assessment with recommendations
    
    Recommended for production use.
    """
    content = await file.read()
    filename = file.filename or "unknown"
    
    file_type = detect_file_type(content, filename)
    text = extract_text_intelligent(content, filename)
    
    if len(text) < 100:
        return {
            "error": f"Insufficient text extracted ({len(text)} chars)",
            "file_type": file_type,
            "suggestion": "Check file format"
        }
    
    result = await smart_parse(text, filename)
    
    return result


@app.get("/health/patterns")
async def get_pattern_health():
    """
    Returns information about pattern coverage and health.
    Use this to understand which patterns exist and their coverage.
    """
    patterns = {
        "total_patterns": 15,
        "patterns": [
            {"id": 1, "name": "Standard date-range", "example": "Company Name | Jan 2020 - Present"},
            {"id": 2, "name": "Worked-as format", "example": "Worked as Developer in Google from 2020 to 2023"},
            {"id": 3, "name": "Table format", "example": "Client: X | Duration: Y"},
            {"id": 4, "name": "Title-DateRange-Client", "example": "Developer (2020-2023) - Client - Employer"},
            {"id": 5, "name": "Standard chronological", "example": "2020-2023: Company Name"},
            {"id": 6, "name": "Bullet format", "example": "• Company: X | Period: 2020-2023"},
            {"id": 7, "name": "ROLE/DESIGNATION keyword", "example": "ROLE: Developer"},
            {"id": 8, "name": "Client with double-dash", "example": "**Client: Company -- Location Date**"},
            {"id": 9, "name": "MM/YYYY format", "example": "01/2020 - 12/2023"},
            {"id": 10, "name": "Markdown format", "example": "## Title | Company | Date"},
            {"id": 11, "name": "Project/Client/Duration", "example": "Project: X | Client: Y | Duration: Z"},
            {"id": 12, "name": "Client/Role format", "example": "Client: X | Role: Y"},
            {"id": 13, "name": "Pipe-separated", "example": "Jun 2015 – Present | Company | Title"},
            {"id": 14, "name": "Tab-separated", "example": "Company, Location [tabs] Date"},
            {"id": 15, "name": "Client with tabs", "example": "Client: Company [tabs] Date"}
        ],
        "coverage": {
            "consultant_resumes": ["11", "12", "15"],
            "corporate_resumes": ["1", "2", "5", "9"],
            "structured_resumes": ["3", "8", "13", "14"],
            "markdown_resumes": ["10"]
        }
    }
    return patterns


@app.post("/parse/agentic")
async def parse_resume_agentic(
    file: UploadFile = File(...),
    use_ai: bool = True
):
    """
    Agentic resume parsing endpoint.
    Uses the self-healing agentic system for robust parsing of any format.
    
    Features:
    - Format auto-detection
    - Gap analysis
    - Targeted AI fixes
    - Failure logging
    """
    content = await file.read()
    filename = file.filename or "unknown"
    
    # Detect file type and extract text
    file_type = detect_file_type(content, filename)
    text = extract_text_intelligent(content, filename)
    
    if len(text) < 100:
        return {
            "error": f"Insufficient text extracted ({len(text)} chars)",
            "file_type": file_type,
            "suggestion": "Check file format or use /debug/diagnosis"
        }
    
    # Run agentic parsing
    result = await agentic_parse(text, filename, use_ai=use_ai)
    
    return result
