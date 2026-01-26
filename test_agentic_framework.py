"""
AGENTIC FRAMEWORK TEST SUITE
============================
Tests the multi-agent resume parsing system:
1. Parser Agent (regex-based)
2. Validation Agent (quality checks)  
3. AI Fallback Agent (Claude API for failures)
"""

import asyncio
import json
import pdfplumber
from docx import Document
from resume_parser_mcp import (
    parse_resume_full, 
    ParseResumeInput,
    validation_agent,
    ai_fallback_agent
)

TEST_RESUMES = [
    # High-quality parses (score >= 80)
    ("Jimmy", "/mnt/user-data/uploads/JimmySunny_ProjectManager.docx", "docx"),
    ("Khaliq", "/mnt/user-data/uploads/Khaliq_IT_PM_Scrum_Master-_Resume_22-SEP-2025.pdf", "pdf"),
    ("Madhuri", "/mnt/user-data/uploads/Madhuri_Mangutkar.docx", "docx"),
    ("Sudheer", "/mnt/user-data/uploads/Sudheer_Babu_Kandukuru_-_Resume.pdf", "pdf"),
    ("Naveen", "/mnt/user-data/uploads/NAVEEN_REDDY_YADLA_Sr_ETL_DEVELOPER.docx", "docx"),
    # Need AI enhancement (score < 80)
    ("Ramaswamy", "/mnt/user-data/uploads/RamaswamyTati-Snowflake.docx", "docx"),
    ("Nageswara", "/mnt/user-data/uploads/Nageswara_Rao_Pamujula_GCP_Data_Engineer2.docx", "docx"),
]

async def run_agentic_test():
    print("=" * 80)
    print("🤖 AGENTIC FRAMEWORK TEST SUITE")
    print("=" * 80)
    print("\nAgents in Pipeline:")
    print("  1. 📝 Parser Agent (Regex-based extraction)")
    print("  2. ✅ Validation Agent (Quality scoring)")
    print("  3. 🤖 AI Fallback Agent (Claude API for failures)")
    print()
    
    results = {"passed": 0, "needs_ai": 0, "total": len(TEST_RESUMES)}
    
    for name, path, fmt in TEST_RESUMES:
        # Load resume
        if fmt == "docx":
            doc = Document(path)
            text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([c.text.strip() for c in row.cells if c.text.strip()])
                    if row_text:
                        text += '\n' + row_text
        else:
            with pdfplumber.open(path) as pdf:
                text = '\n'.join([page.extract_text() or '' for page in pdf.pages])
        
        # Run parser
        result = await parse_resume_full(ParseResumeInput(resume_text=text, use_ai_validation=False))
        data = json.loads(result)
        pr = data["parsed_resume"]
        
        score = data.get("validation_score", 0)
        issues = data.get("validation_issues", [])
        total_resp = sum(len(e.get("responsibilities", [])) for e in pr["experience"])
        
        # Determine status
        if score >= 80:
            status = "✅ PASS"
            results["passed"] += 1
        elif score >= 60:
            status = "⚠️ PARTIAL"
            results["needs_ai"] += 1
        else:
            status = "❌ NEEDS AI"
            results["needs_ai"] += 1
        
        print(f"\n{status} | {name.upper()} | Score: {score}")
        print(f"   Name: {pr['name']} | Email: {pr['email']}")
        print(f"   Exp: {len(pr['experience'])} | Resp: {total_resp} | Edu: {len(pr['education'])}")
        if issues:
            print(f"   Issues: {issues[:3]}{'...' if len(issues) > 3 else ''}")
    
    print("\n" + "=" * 80)
    print("📊 AGENTIC FRAMEWORK RESULTS")
    print("=" * 80)
    print(f"✅ Parser Agent Success: {results['passed']}/{results['total']}")
    print(f"🤖 Need AI Fallback: {results['needs_ai']}/{results['total']}")
    print()
    print("💡 To enable AI Fallback Agent:")
    print("   export ANTHROPIC_API_KEY=your-api-key")
    print("   Then set use_ai_validation=True in ParseResumeInput")
    print("=" * 80)

if __name__ == "__main__":
    asyncio.run(run_agentic_test())
