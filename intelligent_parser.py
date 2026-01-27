"""
Intelligent Resume Parser v3.0 - AI-First Approach
===================================================
Uses Claude AI as the primary extraction engine with regex as validation.
Handles ANY resume format without hardcoded patterns.

Architecture:
1. AI Primary Agent: Claude extracts all structured data
2. Regex Validator: Validates and cleans AI output
3. Quality Agent: Scores and ensures completeness
"""

import json
import re
import os
import asyncio
from datetime import datetime
from dateutil.relativedelta import relativedelta
from typing import Optional, List, Dict, Any, Tuple
from dataclasses import dataclass, field
import unicodedata

# Configuration
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")

# ============================================================================
# TEXT PREPROCESSING
# ============================================================================

def normalize_text(text: str) -> str:
    """Clean and normalize resume text."""
    if not text:
        return ""
    
    # Fix common encoding issues
    replacements = {
        'â€"': '–', 'â€™': "'", 'â€œ': '"', 'â€': '"',
        'Ã©': 'é', 'Ã¨': 'è', 'Ã ': 'à',
        '\u2013': '-', '\u2014': '-', '–': '-',
        '\u2019': "'", '\u2018': "'",
        '\u201c': '"', '\u201d': '"',
        '\u2022': '•', '\u00a0': ' ',
        '\r\n': '\n', '\r': '\n', '\t': ' '
    }
    
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    text = unicodedata.normalize('NFKC', text)
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using multiple methods."""
    text = ""
    
    # Try pdfplumber first (better for complex layouts)
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            text = '\n'.join([p.extract_text() or '' for p in pdf.pages])
    except:
        pass
    
    # Fallback to PyPDF2
    if len(text.strip()) < 100:
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text = '\n'.join([p.extract_text() or '' for p in reader.pages])
        except:
            pass
    
    return normalize_text(text)


def extract_text_from_docx(file_path: str) -> str:
    """Extract ALL text from DOCX: paragraphs + tables + text boxes."""
    from docx import Document
    doc = Document(file_path)
    
    all_text = []
    
    # 1. Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            all_text.append(text)
    
    # 2. Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                all_text.append(' | '.join(row_text))
    
    # 3. Text boxes (for multi-column layouts)
    try:
        for txbx in doc.element.iter():
            if txbx.tag.endswith('txbxContent'):
                texts = [t.text for t in txbx.iter() if t.tag.endswith('}t') and t.text]
                if texts:
                    content = ' '.join(texts)
                    if content not in all_text:
                        all_text.append(content)
    except:
        pass
    
    return normalize_text('\n'.join(all_text))


# ============================================================================
# SIMPLE REGEX EXTRACTORS (High confidence only)
# ============================================================================

def extract_email(text: str) -> Optional[str]:
    """Extract email with high confidence."""
    match = re.search(r'\b([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\b', text)
    return match.group(1).lower() if match else None


def extract_phone(text: str) -> Optional[str]:
    """Extract phone number with various formats."""
    patterns = [
        r'(\+1\s*[-.]?\s*\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4})',
        r'(\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4})',
        r'(\+\d{1,3}[-.\s]?\d{3,4}[-.\s]?\d{3,4}[-.\s]?\d{3,4})',
        r'(\d{10})',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            phone = re.sub(r'[^\d+\-() ]', '', match.group(1)).strip()
            if len(re.sub(r'\D', '', phone)) >= 10:
                return phone
    return None


def extract_linkedin(text: str) -> Optional[str]:
    """Extract LinkedIn URL."""
    patterns = [
        r'linkedin\.com\s*/?\s*in\s*/?\s*([\w-]+)',
        r'linkedin[:\s]+(?:www\.)?linkedin\.com\s*/in/([\w-]+)',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            username = match.group(1)
            if username.lower() not in ['summary', 'profile', 'in']:
                return f"www.linkedin.com/in/{username}"
    return None


# ============================================================================
# AI PRIMARY EXTRACTION
# ============================================================================

AI_EXTRACTION_PROMPT = """You are an expert resume parser. Extract ALL information from this resume into structured JSON.

CRITICAL INSTRUCTIONS:
1. Extract EVERY job/position mentioned, regardless of format
2. Calculate duration_months accurately: (end_year - start_year) * 12 + (end_month - start_month) + 1
3. Extract ALL responsibilities/bullet points for each job
4. Identify the professional title from the most recent role or summary
5. Handle ANY date format (YYYY, MM/YYYY, Month YYYY, MonYY, etc.)
6. For "Present", "Current", or ongoing roles, use end date: {today}
7. Parse client/account information if mentioned (e.g., "Client: ABC Corp" or "UPS Account")

RESUME TEXT:
{resume_text}

Return ONLY valid JSON (no markdown, no explanation, no code blocks) in this exact structure:
{{
  "name": {{
    "full": "Full Name",
    "first": "First",
    "middle": "Middle or null",
    "last": "Last"
  }},
  "contact": {{
    "email": "email@domain.com or null",
    "phone": "phone number or null",
    "linkedin": "linkedin url or null",
    "location": "City, State/Country or null"
  }},
  "professional_title": "Most appropriate title based on experience and summary",
  "summary": "Professional summary if present, or null",
  "experience": [
    {{
      "employer": "Company Name (not client)",
      "title": "Job Title",
      "location": "City, State or null",
      "start_date": "YYYY-MM",
      "end_date": "YYYY-MM",
      "duration_months": 24,
      "is_current": false,
      "client": "Client/Account name if mentioned, or null",
      "responsibilities": [
        "Responsibility 1",
        "Responsibility 2"
      ]
    }}
  ],
  "education": [
    {{
      "degree": "Full Degree Name (e.g., Master of Business Administration, Bachelor of Technology)",
      "field": "Field of Study or null",
      "institution": "University/College Name",
      "year": "YYYY or null",
      "location": "Location or null"
    }}
  ],
  "certifications": [
    "Certification 1 - Issuing Organization",
    "Certification 2"
  ],
  "skills": {{
    "technical": ["Python", "Java", "SQL"],
    "tools": ["AWS", "Docker", "Kubernetes"],
    "databases": ["PostgreSQL", "MongoDB"],
    "frameworks": ["React", "Django"]
  }}
}}

IMPORTANT RULES:
- Sort experience by start_date descending (most recent first)
- Include ALL jobs, even short-term or contract positions
- Extract complete responsibilities, not truncated
- For education without explicit year, extract from context if possible
- Clean up any encoding artifacts in names/text
- If month is not specified, use 01 for start dates and 12 for end dates
"""


async def ai_extract_resume(text: str) -> Dict:
    """Use Claude AI to extract all resume data."""
    if not ANTHROPIC_API_KEY:
        return {"error": "No API key configured"}
    
    try:
        import httpx
        
        today = datetime.now().strftime("%Y-%m")
        prompt = AI_EXTRACTION_PROMPT.format(
            today=today,
            resume_text=text[:25000]  # Limit to avoid token issues
        )
        
        async with httpx.AsyncClient(timeout=120.0) as client:
            response = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": ANTHROPIC_API_KEY,
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json"
                },
                json={
                    "model": "claude-sonnet-4-20250514",
                    "max_tokens": 8000,
                    "messages": [{"role": "user", "content": prompt}]
                }
            )
            
            if response.status_code == 200:
                result = response.json()
                ai_text = result.get("content", [{}])[0].get("text", "")
                
                # Extract JSON from response (handle potential markdown)
                ai_text = ai_text.strip()
                if ai_text.startswith('```'):
                    ai_text = re.sub(r'^```(?:json)?\n?', '', ai_text)
                    ai_text = re.sub(r'\n?```$', '', ai_text)
                
                json_match = re.search(r'\{[\s\S]*\}', ai_text)
                if json_match:
                    return json.loads(json_match.group())
                return {"error": "No JSON found in response"}
            else:
                return {"error": f"API returned {response.status_code}: {response.text[:200]}"}
                
    except json.JSONDecodeError as e:
        return {"error": f"JSON parse error: {str(e)}"}
    except Exception as e:
        return {"error": str(e)}
    
    return {"error": "Unknown error"}


# ============================================================================
# QUALITY VALIDATION & SCORING
# ============================================================================

def validate_and_score(parsed: Dict, original_text: str) -> Tuple[Dict, int, List[str]]:
    """Validate AI output and calculate quality score."""
    issues = []
    score = 100
    
    # Validate name
    name = parsed.get("name", {})
    if not name.get("full") or len(name.get("full", "")) < 3:
        issues.append("missing_name")
        score -= 20
    
    # Validate contact (email is critical)
    contact = parsed.get("contact", {})
    if not contact.get("email"):
        # Try regex fallback
        email = extract_email(original_text)
        if email:
            contact["email"] = email
        else:
            issues.append("missing_email")
            score -= 15
    
    # Validate phone
    if not contact.get("phone"):
        phone = extract_phone(original_text)
        if phone:
            contact["phone"] = phone
    
    # Validate LinkedIn
    if not contact.get("linkedin"):
        linkedin = extract_linkedin(original_text)
        if linkedin:
            contact["linkedin"] = linkedin
    
    # Validate experience
    experience = parsed.get("experience", [])
    if not experience:
        issues.append("missing_experience")
        score -= 25
    else:
        total_resp = sum(len(e.get("responsibilities", [])) for e in experience)
        if total_resp < 3:
            issues.append("low_responsibilities")
            score -= 10
        
        # Check and fix each experience entry
        for i, exp in enumerate(experience):
            if not exp.get("employer"):
                issues.append(f"exp_{i}_missing_employer")
                score -= 5
            if not exp.get("title"):
                issues.append(f"exp_{i}_missing_title")
                score -= 5
            
            # Ensure duration is calculated
            if not exp.get("duration_months") or exp.get("duration_months", 0) <= 0:
                exp["duration_months"] = calculate_duration_from_dates(
                    exp.get("start_date"), exp.get("end_date")
                )
    
    # Validate professional title
    if not parsed.get("professional_title"):
        if experience:
            parsed["professional_title"] = experience[0].get("title", "Professional")
        else:
            issues.append("missing_title")
            score -= 10
    
    # Validate education
    education = parsed.get("education", [])
    if not education:
        issues.append("missing_education")
        score -= 10
    
    # Calculate totals
    total_months = sum(e.get("duration_months", 0) for e in experience)
    parsed["total_experience_months"] = total_months
    parsed["total_experience_years"] = round(total_months / 12, 1) if total_months else 0
    
    return parsed, max(0, score), issues


def calculate_duration_from_dates(start: str, end: str) -> int:
    """Calculate duration in months from date strings."""
    if not start:
        return 0
    
    try:
        # Parse start date
        start_parts = start.split('-')
        start_year = int(start_parts[0])
        start_month = int(start_parts[1]) if len(start_parts) > 1 else 1
        
        # Parse end date
        if not end or end.lower() in ['present', 'current']:
            end_year = datetime.now().year
            end_month = datetime.now().month
        else:
            end_parts = end.split('-')
            end_year = int(end_parts[0])
            end_month = int(end_parts[1]) if len(end_parts) > 1 else 12
        
        # Calculate duration
        months = (end_year - start_year) * 12 + (end_month - start_month) + 1
        return max(1, months)
    except:
        return 0


# ============================================================================
# OUTPUT FORMATTING
# ============================================================================

def format_output(parsed: Dict) -> Dict:
    """Format parsed data into final output structure."""
    name = parsed.get("name", {})
    contact = parsed.get("contact", {})
    
    # Build experience list
    experience_list = []
    for exp in parsed.get("experience", []):
        experience_list.append({
            "Employer": exp.get("employer"),
            "title": exp.get("title"),
            "location": exp.get("location"),
            "start_date": exp.get("start_date"),
            "end_date": exp.get("end_date"),
            "duration_months": exp.get("duration_months", 0),
            "responsibilities": exp.get("responsibilities", []),
            "tools": [],
            "client": exp.get("client")
        })
    
    # Build education list
    education_list = []
    for edu in parsed.get("education", []):
        degree = edu.get("degree", "")
        if edu.get("field") and edu.get("field") not in str(degree):
            degree = f"{degree} in {edu['field']}" if degree else edu['field']
        
        education_list.append({
            "degree": degree or None,
            "institution": edu.get("institution"),
            "year": edu.get("year")
        })
    
    # Build skills list
    skills = parsed.get("skills", {})
    technical_skills = []
    for category in ['technical', 'tools', 'databases', 'frameworks', 'languages']:
        technical_skills.extend(skills.get(category, []))
    
    return {
        "parsed_resume": {
            "firstname": name.get("first"),
            "lastname": name.get("last"),
            "name": name.get("full"),
            "title": parsed.get("professional_title"),
            "location": contact.get("location"),
            "phone_number": contact.get("phone"),
            "email": contact.get("email"),
            "linkedin": contact.get("linkedin"),
            "summary": parsed.get("summary"),
            "total_experience_months": parsed.get("total_experience_months", 0),
            "total_experience_years": parsed.get("total_experience_years", 0),
            "technical_skills": list(set(technical_skills)),
            "education": education_list,
            "certifications": parsed.get("certifications", []),
            "experience": experience_list
        }
    }


# ============================================================================
# MAIN PARSING FUNCTION
# ============================================================================

async def parse_resume_intelligent(
    text: str,
    filename: str = "",
    file_path: str = None
) -> str:
    """
    Intelligent resume parser using AI-first approach.
    
    Args:
        text: Resume text content
        filename: Original filename
        file_path: Path to file for enhanced extraction
    
    Returns:
        JSON string with parsed resume data
    """
    # Normalize text
    text = normalize_text(text)
    
    # If file path provided, try enhanced extraction
    if file_path:
        if file_path.endswith('.pdf'):
            enhanced_text = extract_text_from_pdf(file_path)
            if len(enhanced_text) > len(text):
                text = enhanced_text
        elif file_path.endswith('.docx'):
            enhanced_text = extract_text_from_docx(file_path)
            if len(enhanced_text) > len(text):
                text = enhanced_text
    
    # Step 1: AI Primary Extraction
    ai_result = await ai_extract_resume(text)
    
    if "error" in ai_result:
        # AI failed - return error with basic regex extraction
        return json.dumps({
            "parsed_resume": {
                "name": None,
                "email": extract_email(text),
                "phone_number": extract_phone(text),
                "linkedin": extract_linkedin(text),
                "experience": [],
                "education": [],
                "filename": filename
            },
            "ai_error": ai_result["error"],
            "validation_score": 0,
            "validation_issues": ["ai_extraction_failed"]
        }, indent=2)
    
    # Step 2: Validate and Score
    validated, score, issues = validate_and_score(ai_result, text)
    
    # Step 3: Format Output
    output = format_output(validated)
    output["parsed_resume"]["filename"] = filename
    output["validation_score"] = score
    output["validation_issues"] = issues
    output["ai_enhanced"] = True
    
    return json.dumps(output, indent=2, ensure_ascii=False)


# ============================================================================
# BACKWARD COMPATIBLE INTERFACE
# ============================================================================

# Pydantic models for compatibility
try:
    from pydantic import BaseModel, Field, ConfigDict
    from enum import Enum
    
    class ResponseFormat(str, Enum):
        MARKDOWN = "markdown"
        JSON = "json"
    
    class ParseResumeInput(BaseModel):
        model_config = ConfigDict(str_strip_whitespace=True)
        resume_text: str = Field(..., min_length=50, max_length=100000)
        response_format: ResponseFormat = Field(default=ResponseFormat.JSON)
        filename: Optional[str] = Field(default=None)
        file_path: Optional[str] = Field(default=None)
        use_ai_validation: bool = Field(default=True)

except ImportError:
    pass


async def parse_resume_full(params) -> str:
    """Backward compatible interface for existing code."""
    return await parse_resume_intelligent(
        text=params.resume_text,
        filename=params.filename or "",
        file_path=params.file_path
    )


# ============================================================================
# CLI TESTING
# ============================================================================

if __name__ == "__main__":
    import sys
    
    async def test_file(path: str):
        if path.endswith('.pdf'):
            text = extract_text_from_pdf(path)
        elif path.endswith('.docx'):
            text = extract_text_from_docx(path)
        else:
            with open(path, 'r') as f:
                text = f.read()
        
        result = await parse_resume_intelligent(text, path.split('/')[-1], path)
        print(result)
    
    if len(sys.argv) > 1:
        asyncio.run(test_file(sys.argv[1]))
    else:
        print("Usage: python intelligent_parser.py <resume_file>")
